"""Offline rule-coverage regression test.

Proves the validation ENGINE works, two ways, with no SharePoint / Azure /
network:

  1. DETECTION  — with auto_fix off, every implemented Word check reports its
     violation.
  2. AUTO-FIX   — with auto_fix on, every fixable check actually rewrites the
     document (text replaced, or font/colour corrected), and the detect-only
     checks still flag without claiming a fix.

Why this exists: the previous local test (test_local.py) only exercised 2 of
the 6 rule types with idealised data, so it passed while real documents failed.
This FAILS LOUDLY (exit 1) if any rule stops catching or stops fixing its
violation — a regression gate for the engine.

Run:  python3 test_rule_coverage.py
Pair with rule_doctor.py, which checks your REAL SharePoint rules are shaped to
actually reach this engine.
"""
import sys
from io import BytesIO

from docx import Document
from docx.shared import RGBColor

from ValidateDocument.word_validator import validate_word_document
from ValidateDocument.enhanced_validators import iter_all_paragraphs

MACE_BLUE = (0, 51, 153)

# Text-based checks. Each carries the document text and, where the check is
# auto-fixable, what the corrected text must / must not contain. detect_only
# checks report a violation but have no fix branch in the validators.
TEXT_CASES = [
    {"title": "British spelling (color->colour)", "rule_type": "Language",
     "check_value": "BritishSpelling_color", "expected_value": "colour",
     "text": "The color of the cladding was approved.",
     "fix_present": "colour", "fix_absent": "color "},
    {"title": "No contraction (don't)", "rule_type": "Language",
     "check_value": "NoContraction_dont", "expected_value": "do not",
     "text": "We don't accept late submissions.",
     "fix_present": "do not", "fix_absent": "don't"},
    {"title": "Word choice (towards->toward)", "rule_type": "Language",
     "check_value": "Word_toward", "expected_value": "toward",
     "text": "Progress was made towards completion.",
     "fix_present": "toward", "fix_absent": "towards"},
    {"title": "Avoid etc.", "rule_type": "Language",
     "check_value": "AvoidEtc", "expected_value": "",
     "text": "Bring drawings, specifications, etc. to the meeting.",
     "detect_only": True},
    {"title": "No ampersand (&)", "rule_type": "Punctuation",
     "check_value": "NoAmpersand", "expected_value": "and",
     "text": "Design & build was selected.",
     "fix_present": "and", "fix_absent": "&"},
    {"title": "Percent symbol (%)", "rule_type": "Punctuation",
     "check_value": "PercentSymbol", "expected_value": "percent",
     "text": "The works are 85% complete.",
     "fix_present": "85 percent", "fix_absent": "%"},
    {"title": "No apostrophe plurals (CD's)", "rule_type": "Punctuation",
     "check_value": "NoApostrophePlurals", "expected_value": "",
     "text": "Several CD's were issued to the team.",
     "detect_only": True},
    {"title": "Number commas (1000000)", "rule_type": "Punctuation",
     "check_value": "NumberCommas", "expected_value": "",
     "text": "The budget is 1000000 pounds for this phase.",
     "fix_present": "1,000,000", "fix_absent": "1000000"},
]

# Font/colour checks are asserted by inspecting runs, not text.
FONT_SENTENCE = "This sentence is in Times New Roman."
ALLTEXTFONT = {"title": "All text font (Arial)", "rule_type": "Font",
               "check_value": "AllTextFont", "expected_value": "Arial"}
HEADING1FONT = {"title": "Heading 1 font (Arial)", "rule_type": "Font",
                "check_value": "Heading1Font", "expected_value": "Arial"}
HEADING1COLOR = {"title": "Heading 1 colour (Mace blue)", "rule_type": "Color",
                 "check_value": "Heading1Color", "expected_value": "0,51,153"}


def _rule(case, auto_fix):
    return {"title": case["title"], "rule_type": case["rule_type"], "doc_type": "Word",
            "check_value": case["check_value"], "expected_value": case["expected_value"],
            "auto_fix": auto_fix, "use_ai": False, "priority": 10}


def _build_document():
    doc = Document()
    heading = doc.add_heading("Project Report", level=1)
    if heading.runs:
        heading.runs[0].font.name = "Times New Roman"
        heading.runs[0].font.color.rgb = RGBColor(255, 0, 0)
    for case in TEXT_CASES:
        doc.add_paragraph(case["text"])
    p = doc.add_paragraph()
    p.add_run(FONT_SENTENCE).font.name = "Times New Roman"
    stream = BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream


def _all_text(doc):
    return "\n".join(p.text for p in iter_all_paragraphs(doc))


def _all_rules(auto_fix):
    rules = [_rule(c, auto_fix) for c in TEXT_CASES]
    rules += [_rule(c, auto_fix) for c in (ALLTEXTFONT, HEADING1FONT, HEADING1COLOR)]
    return rules


def _line(result, label):
    print(f"  {'PASS' if result else 'FAIL':6}  {label}")
    return result


def test_detection():
    """auto_fix off: every check must report its violation."""
    print("\n[1] Detection — every implemented check flags its violation\n")
    result = validate_word_document(_build_document(), _all_rules(auto_fix=False))
    fired = {i.get("rule_name") for i in result["issues"]}
    fired |= {f.get("rule_name") for f in result["fixes_applied"]}
    cases = TEXT_CASES + [ALLTEXTFONT, HEADING1FONT, HEADING1COLOR]
    return all(_line(c["title"] in fired, f"{c['check_value']:24}  {c['title']}") for c in cases)


def test_autofix():
    """auto_fix on: fixable checks rewrite the doc; detect-only checks still flag."""
    print("\n[2] Auto-fix — fixable checks actually correct the document\n")
    result = validate_word_document(_build_document(), _all_rules(auto_fix=True))
    doc = result["document"]
    text = _all_text(doc)
    fired = {i.get("rule_name") for i in result["issues"]}
    fired |= {f.get("rule_name") for f in result["fixes_applied"]}
    oks = []

    for case in TEXT_CASES:
        if case.get("detect_only"):
            ok = case["title"] in fired
            oks.append(_line(ok, f"{case['check_value']:24}  {case['title']} (detect-only, still flagged)"))
            continue
        present = case["fix_present"] in text
        absent = case["fix_absent"] not in text
        oks.append(_line(present and absent,
                         f"{case['check_value']:24}  {case['title']} "
                         f"-> '{case['fix_present']}'" + ("" if absent else f" [still has '{case['fix_absent']}']")))

    # Font: the Times New Roman body run must now be Arial.
    body_fonts = [r.font.name for p in iter_all_paragraphs(doc) for r in p.runs
                  if r.text and FONT_SENTENCE[:12] in r.text]
    oks.append(_line(bool(body_fonts) and all(f == "Arial" for f in body_fonts),
                     f"{'AllTextFont':24}  body run corrected to Arial"))

    # Heading font + colour.
    h_runs = [r for p in iter_all_paragraphs(doc) if p.style.name == "Heading 1" for r in p.runs]
    oks.append(_line(bool(h_runs) and all(r.font.name == "Arial" for r in h_runs),
                     f"{'Heading1Font':24}  heading corrected to Arial"))
    oks.append(_line(bool(h_runs) and all(r.font.color.rgb == RGBColor(*MACE_BLUE) for r in h_runs if r.font.color and r.font.color.rgb),
                     f"{'Heading1Color':24}  heading corrected to Mace blue"))
    return all(oks)


# Deterministic batch-2 checks. title == check_value. Mostly detection-only;
# two auto-fix (NoSpacesAroundSlash, Hyphen_wide).
NEW_CASES = [
    {"rt": "Punctuation", "cv": "TimeFormat", "text": "The site opens at 9am each day."},
    {"rt": "Punctuation", "cv": "DateFormat_Text", "text": "Issued on 01/02/2015 for review."},
    {"rt": "Punctuation", "cv": "YearIntervalFormat", "text": "The 2019-2020 programme is approved."},
    {"rt": "Punctuation", "cv": "NoSpacesAroundSlash", "text": "Velocity measured in km / s here.",
     "fix_present": "km/s", "fix_absent": " / "},
    {"rt": "Punctuation", "cv": "AvoidForwardSlash", "text": "Submit to client/contractor today."},
    {"rt": "Punctuation", "cv": "HyphenInWords", "text": "An in depth review is required."},
    {"rt": "Punctuation", "cv": "HyphenSuffixes", "text": "A quality related issue arose."},
    {"rt": "Punctuation", "cv": "HyphenAlwaysPrefix", "text": "He is a self made leader."},
    {"rt": "Punctuation", "cv": "Hyphen_wide", "text": "A site wide policy applies now.",
     "fix_present": "site-wide", "fix_absent": "site wide"},
    {"rt": "Punctuation", "cv": "PunctuationBeforeEgIe", "text": "Use tools e.g. drawings for this."},
    {"rt": "Punctuation", "cv": "OxfordComma", "text": "We need bricks, mortar and sand here."},
    {"rt": "Punctuation", "cv": "NumbersBelowTen", "text": "There are 3 teams on the site."},
    {"rt": "Punctuation", "cv": "CompoundModifiers", "text": "Submit the 15 page report today."},
    {"rt": "Punctuation", "cv": "CaptionNoPeriod", "text": "Figure 1 The site location plan.",
     "caption": True},
    {"rt": "Grammar", "cv": "NoSentenceStartEgIe", "text": "E.g. the north wing is complete."},
    {"rt": "Grammar", "cv": "NoEtcWithEgIe", "text": "Bring tools e.g. drills, etc. today."},
    {"rt": "Grammar", "cv": "ClientNameNotTheClient", "text": "Send it to the client for sign-off."},
    {"rt": "Grammar", "cv": "OrgSingular", "text": "The team are reviewing it now."},
    {"rt": "Language", "cv": "NoFeelTechnical", "text": "We feel the design is adequate."},
    {"rt": "Language", "cv": "NoAboveBelow", "text": "See above for the full details."},
    {"rt": "Language", "cv": "PreferMetric", "text": "The access route is 5 miles long."},
    {"rt": "Capitalisation", "cv": "ProperNounDerivations", "text": "The welsh survey was completed."},
    {"rt": "Capitalisation", "cv": "NoEmphasisCaps", "text": "This is VERY IMPORTANT indeed."},
]


def _new_rule(case, auto_fix):
    return {"title": case["cv"], "rule_type": case["rt"], "doc_type": "Word",
            "check_value": case["cv"], "expected_value": "", "auto_fix": auto_fix,
            "use_ai": False, "priority": 10}


def _build_new_doc():
    doc = Document()
    for case in NEW_CASES:
        if case.get("caption"):
            try:
                doc.add_paragraph(case["text"], style="Caption")
            except KeyError:
                doc.add_paragraph(case["text"])  # template lacks Caption style
        else:
            doc.add_paragraph(case["text"])
    stream = BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream


def test_new_checks():
    """Batch-2 deterministic checks: detection for all, rewrite for the fixable two."""
    print("\n[3] Deterministic batch-2 — newly implemented checks fire\n")
    det = validate_word_document(_build_new_doc(), [_new_rule(c, False) for c in NEW_CASES])
    fired = {i.get("rule_name") for i in det["issues"]} | {f.get("rule_name") for f in det["fixes_applied"]}
    oks = [_line(c["cv"] in fired, f"{c['cv']:24}  {c['text']}") for c in NEW_CASES]

    # Auto-fix rewrite for the two fixable checks.
    fix = validate_word_document(_build_new_doc(), [_new_rule(c, True) for c in NEW_CASES])
    text = _all_text(fix["document"])
    for c in NEW_CASES:
        if "fix_present" in c:
            ok = c["fix_present"] in text and c["fix_absent"] not in text
            oks.append(_line(ok, f"{c['cv']:24}  rewrite -> '{c['fix_present']}'"))
    return all(oks)


def run():
    detection_ok = test_detection()
    autofix_ok = test_autofix()
    new_ok = test_new_checks()
    print()
    if detection_ok and autofix_ok and new_ok:
        print("  ✓ All checks detect AND fix their violations. The engine works end to end;")
        print("    if a real rule isn't firing, run rule_doctor.py on the real rules.")
        return 0
    print("  ✗ Some checks failed above — the engine is not applying/fixing them.")
    print("    Investigate the matching validator in enhanced_validators.py / word_validator.py.")
    return 1


if __name__ == "__main__":
    sys.exit(run())

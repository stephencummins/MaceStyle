"""
Create a test Word document that breaks all style rules
"""
from docx import Document
from docx.shared import Pt, RGBColor

def create_test_document():
    """Create a Word document that violates all the style rules"""

    doc = Document()

    # ============================================
    # RULE VIOLATIONS
    # ============================================

    # Add title with wrong font (Times New Roman instead of Arial)
    title = doc.add_heading('PROJECT MANAGEMENT PLAN', level=1)
    for run in title.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16)

    # Introduction section - multiple violations
    intro = doc.add_heading('1. INTRODUCTION AND BACKGROUND', level=2)
    for run in intro.runs:
        run.font.name = 'Calibri'  # Wrong font

    # Paragraph with multiple violations
    p1 = doc.add_paragraph()

    # Use Calibri font (wrong - should be Arial)
    text1 = p1.add_run(
        "This document can't be finalized until we analyze the color scheme for the new center. "
        "We don't have enough information yet, and it isn't clear when we'll receive it. "
        "The project won't start until 3/15/2025 at 1:30 PM. "
    )
    text1.font.name = 'Calibri'
    text1.font.size = Pt(11)

    # Add more violations
    p2 = doc.add_paragraph()
    text2 = p2.add_run(
        "The Budget is approximately 5000000 dollars & includes 5 separate work packages. "
        "We should organize the team meeting towards the end of the month, etc. "
        "The success rate is around 85% based on our analysis. "
    )
    text2.font.name = 'Times New Roman'  # Wrong font
    text2.font.size = Pt(11)

    # Capitalisation violations
    p3 = doc.add_paragraph()
    text3 = p3.add_run(
        "The Project Manager and the Senior Engineer will coordinate with the Client "
        "to ensure all Requirements are met. The Company has established clear Guidelines "
        "for all Team Members to follow throughout the Project Lifecycle. "
    )
    text3.font.name = 'Verdana'  # Wrong font

    # British vs American spelling violations
    heading2 = doc.add_heading('2. project scope and organization', level=2)  # Wrong capitalization
    for run in heading2.runs:
        run.font.name = 'Georgia'  # Wrong font

    p4 = doc.add_paragraph()
    text4 = p4.add_run(
        "We must prioritize and optimize our aluminum procurement strategy. "
        "The program's organization will be centralized to maximize efficiency. "
        "We'll license the software and analyze the harbor infrastructure. "
    )
    text4.font.name = 'Cambria'  # Wrong font

    # More violations - contractions and word choice
    p5 = doc.add_paragraph()
    text5 = p5.add_run(
        "We couldn't complete the review because we didn't have the data. "
        "It shouldn't take long, but we haven't confirmed the timeline yet. "
        "We could start next week if the team's available. "
    )
    text5.font.name = 'Calibri'  # Wrong font

    # Numbers and hyphenation violations
    p6 = doc.add_paragraph()
    text6 = p6.add_run(
        "The site wide assessment covers 3 different areas with a total of 12500 square meters. "
        "We identified 7 critical issues and 4 minor concerns. "
        "The 15 page report will be submitted by 28-Feb-2025. "  # Wrong date format
    )
    text6.font.name = 'Arial'  # Correct font but other violations

    # Quotation violations
    p7 = doc.add_paragraph()
    text7 = p7.add_run(
        'The term "best practice" should be applied consistently. '  # Should use single quotes
        "The manager said, 'We need to improve our processes.' "  # Should use double quotes
        "The team's responsibilities include quality control, risk management, etc. "
    )
    text7.font.name = 'Tahoma'  # Wrong font

    # Apostrophe violations
    p8 = doc.add_paragraph()
    text8 = p8.add_run(
        "We received multiple CD's and DVD's from the supplier's. "  # Wrong apostrophes
        "All the SME's agreed that the KPI's need improvement. "
    )
    text8.font.name = 'Comic Sans MS'  # Very wrong font!

    # Symbol violations
    heading3 = doc.add_heading('3. Roles & Responsibilities', level=2)  # & instead of 'and'
    for run in heading3.runs:
        run.font.name = 'Impact'  # Wrong font

    p9 = doc.add_paragraph()
    text9 = p9.add_run(
        "The Project Director & Programme Manager will coordinate towards achieving "
        "a 95% completion rate. The team should focus on the following area's: "
    )
    text9.font.name = 'Arial Narrow'  # Wrong font

    # Bulleted list with violations
    doc.add_paragraph('Quality Assurance & Control', style='List Bullet')
    doc.add_paragraph('Health & Safety Management', style='List Bullet')
    doc.add_paragraph('Risk & Issue Management', style='List Bullet')

    # Apply wrong font to list items
    for paragraph in doc.paragraphs[-3:]:
        for run in paragraph.runs:
            run.font.name = 'Courier New'

    # More violations
    p10 = doc.add_paragraph()
    text10 = p10.add_run(
        "The organization couldn't meet it's target's due to unforseen circumstances. "
        "We analysed the situation & realized we should've started earlier. "
        "The utilization of resources wasn't optimized towards the project goal's. "
    )
    text10.font.name = 'Book Antiqua'  # Wrong font

    # Date and time violations
    heading4 = doc.add_heading('4. KEY MILESTONES & DELIVERABLES', level=2)
    for run in heading4.runs:
        run.font.name = 'Palatino Linotype'  # Wrong font

    p11 = doc.add_paragraph()
    text11 = p11.add_run(
        "Project kickoff: 03/01/2025 at 2:00pm\n"  # Wrong date format, wrong time format
        "Design review: 15/03/2025 at 10:30am\n"  # Wrong formats
        "Construction start: April 1st, 2025 at 8:00am\n"  # Inconsistent format
        "Completion target: 12-31-2025\n"  # Wrong format
    )
    text11.font.name = 'Franklin Gothic Medium'  # Wrong font

    # Compound modifier violations (missing hyphens)
    p12 = doc.add_paragraph()
    text12 = p12.add_run(
        "The 20 page technical specification document outlines the programme wide requirements. "
        "Our state of the art facility will support the 5 year development plan. "
        "The high level design covers 8 work packages. "
    )
    text12.font.name = 'Trebuchet MS'  # Wrong font

    # More American spellings
    p13 = doc.add_paragraph()
    text13 = p13.add_run(
        "We need to finalize the labor costs and catalog all specialized equipment. "
        "The defense strategy prioritizes minimizing risk through rigorous modeling. "
        "We'll utilize fiber optic cables in the theater of operations. "
    )
    text13.font.name = 'Garamond'  # Wrong font

    # Word choice violations
    heading5 = doc.add_heading('5. conclusion', level=2)  # Wrong capitalization
    for run in heading5.runs:
        run.font.name = 'Century Gothic'  # Wrong font

    p14 = doc.add_paragraph()
    text14 = p14.add_run(
        "In the event of any issues, we should contact the Project Manager. "
        "For the purpose of maintaining quality, we could implement additional checks. "
        "At the present time, we're on schedule despite the fact that we've had some delays. "
        "In order to complete on time, we'll need 3 to 5 additional resources. "
    )
    text14.font.name = 'Lucida Sans'  # Wrong font

    # Final paragraph with mixed violations
    p15 = doc.add_paragraph()
    text15 = p15.add_run(
        "This Programme can't proceed until we've finalized the organisation's requirements. "
        "The Centre of Excellence should've provided guidance towards optimizing our approach. "
        "We identified 2500 line item's totaling 15500000 dollars & we analysed each one carefully. "
        "The teams performance exceeded expectation's by 12%. "
    )
    text15.font.name = 'Arial Black'  # Wrong font (Arial Black not Arial)

    # Save the document
    output_path = '/Users/stephen/Desktop/test_document_with_violations.docx'
    doc.save(output_path)

    print(f"✅ Test document created: {output_path}")
    print(f"\n📊 Violations included:")
    print(f"   • Multiple font violations (should all be Arial)")
    print(f"   • American spellings: color, center, analyze, organization, aluminum, license, harbor, etc.")
    print(f"   • Contractions: can't, don't, isn't, won't, couldn't, didn't, shouldn't, haven't, could've, should've, etc.")
    print(f"   • Wrong date formats: 3/15/2025, 28-Feb-2025, 03/01/2025, etc.")
    print(f"   • Wrong time formats: 1:30 PM, 2:00pm, 10:30am, etc.")
    print(f"   • Numbers below 10 not spelled out: 3, 5, 7, 4, 8, etc.")
    print(f"   • Missing commas in large numbers: 5000000, 12500, 15500000, 2500")
    print(f"   • Wrong capitalization: PROJECT, INTRODUCTION, etc.")
    print(f"   • Ampersands instead of 'and': &")
    print(f"   • Percent symbol instead of 'percent': %")
    print(f"   • Wrong apostrophes: CD's, DVD's, area's, target's, goal's, it's")
    print(f"   • Missing hyphens: site wide, 20 page, programme wide, 5 year, state of the art")
    print(f"   • Wrong quote types")
    print(f"   • Word choice: towards, etc., should/could")
    print(f"\n🧪 This document should trigger many validation failures!")

if __name__ == "__main__":
    create_test_document()

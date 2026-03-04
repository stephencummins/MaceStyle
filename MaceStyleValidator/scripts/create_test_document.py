"""
Create a comprehensive test Word document with all style violations
"""
from docx import Document
from docx.shared import Pt, RGBColor
import os

def create_test_document():
    """Create a test Word document with multiple style violations"""

    print("Creating comprehensive test Word document...\n")

    # Create a new Document
    doc = Document()

    # Add title
    title = doc.add_heading('Mace Style Validator Test Document', 0)

    # Add introduction
    intro = doc.add_paragraph(
        "This document contains intentional style violations to test the Mace Style Validator. "
        "Each section below demonstrates different types of errors that should be detected and corrected."
    )

    # Section 1: British English Spelling Errors
    doc.add_heading('1. British English Spelling Violations', 1)

    violations_british = [
        "The project has been finalized and is ready for review.",
        "We need to analyze the color scheme used in the center of the building.",
        "The organization will authorize the changes after the defense review.",
        "Our labor costs have been optimized using an analog meter.",
        "The fiber optic cables were gray and needed to be organized properly.",
        "The program will maximize efficiency across the harbor.",
        "We traveled to the neighbor's building and realized the true caliber of their work.",
    ]

    for violation in violations_british:
        p = doc.add_paragraph(violation, style='List Bullet')
        # Set some text to wrong font for testing
        if p.runs:
            p.runs[0].font.name = 'Calibri'

    # Section 2: Contractions (Grammar)
    doc.add_heading('2. Contraction Violations', 1)

    violations_contractions = [
        "We can't proceed without proper authorization.",
        "They don't have the required permits yet.",
        "The system won't work if it isn't configured properly.",
        "You shouldn't submit until you've reviewed everything.",
        "We didn't receive the materials, so we couldn't complete the task.",
        "It hasn't been approved, and we haven't scheduled the review.",
        "They're waiting for confirmation that we're ready to proceed.",
    ]

    for violation in violations_contractions:
        p = doc.add_paragraph(violation, style='List Bullet')
        if p.runs:
            p.runs[0].font.name = 'Times New Roman'

    # Section 3: Symbols and Punctuation
    doc.add_heading('3. Symbol & Punctuation Violations', 1)

    violations_symbols = [
        "The partnership between M&S and other retailers showed 50% growth.",
        "Johnson & Johnson reported 75% completion of their facilities.",
        "The R&D department has 100% commitment to quality & safety.",
        "Cost reduction of 25% was achieved in Q&A sessions.",
    ]

    for violation in violations_symbols:
        doc.add_paragraph(violation, style='List Bullet')

    # Section 4: Number Formatting
    doc.add_heading('4. Number Formatting Violations', 1)

    violations_numbers = [
        "The project cost 1000000 pounds and took 2500 hours to complete.",
        "We received 5000 applications from 10000 potential candidates.",
        "The building spans 50000 square feet across 3 floors.",
        "Budget allocation: 15000 for materials, 25000 for labor, 10000 for equipment.",
    ]

    for violation in violations_numbers:
        doc.add_paragraph(violation, style='List Bullet')

    # Section 5: Mixed Violations
    doc.add_heading('5. Combined Violations', 1)

    mixed_violations = [
        "The finalized analysis can't be submitted to the organization until we've received authorization from M&S.",
        "We don't believe the color of the center's harbor meets our defense standards - it's gray & doesn't match our program.",
        "The labor costs won't be optimized until 5000 employees are organized properly in the neighbor's facility.",
        "They're analyzing 1000 fiber optic connections but can't maximize efficiency by 50%.",
    ]

    for violation in mixed_violations:
        doc.add_paragraph(violation, style='List Bullet')

    # Section 6: Font Testing
    doc.add_heading('6. Font Consistency Test', 1)

    p = doc.add_paragraph()
    p.add_run("This text is in Arial (correct). ").font.name = 'Arial'
    p.add_run("This text is in Calibri (wrong). ").font.name = 'Calibri'
    p.add_run("This text is in Times New Roman (wrong). ").font.name = 'Times New Roman'
    p.add_run("This text is in Comic Sans (very wrong). ").font.name = 'Comic Sans MS'
    p.add_run("This text has no font set (should default to Arial).").font.name = None

    # Summary section
    doc.add_heading('Expected Corrections Summary', 1)

    summary_points = [
        "British English: finalized→finalised, color→colour, center→centre, organization→organisation, etc.",
        "Contractions: can't→cannot, don't→do not, won't→will not, isn't→is not, etc.",
        "Symbols: M&S→M and S, &→and, 50%→50 percent, 75%→75 percent",
        "Numbers: 1000→1,000, 5000→5,000, 1000000→1,000,000",
        "Fonts: All text should be Arial",
    ]

    for point in summary_points:
        doc.add_paragraph(point, style='List Bullet')

    # Save the document
    output_dir = "/Users/stephen/Projects/MaceStyle/test_files"
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, "test_validation_comprehensive.docx")

    doc.save(output_path)

    print("✅ Test document created successfully!")
    print(f"\n📄 Location: {output_path}")
    print("\n📊 Document contains:")
    print("   • 7 British English spelling violations")
    print("   • 7 contraction violations")
    print("   • 4 symbol violations (& and %)")
    print("   • 4 number formatting violations")
    print("   • 4 combined/complex violations")
    print("   • Multiple font violations")
    print("\n🎯 Total violations: 40+")
    print("\n📋 Next steps:")
    print("   1. Upload this file to your SharePoint Document Library")
    print("   2. Wait for automatic validation (5-15 seconds)")
    print("   3. Check the validation status and HTML report")
    print("   4. Download the corrected file to verify all fixes")
    print("\n" + "="*70)

    return output_path

if __name__ == "__main__":
    try:
        create_test_document()
    except Exception as e:
        print(f"\n❌ ERROR: {e}")
        import traceback
        traceback.print_exc()

"""
Create a test document with intentional style violations
"""
import sys
sys.path.append('ValidateDocument')

from ValidateDocument import get_graph_token, download_file, upload_file
from docx import Document
from docx.shared import Pt, RGBColor
from io import BytesIO

def main():
    print("📝 Creating test document with style violations\n")

    try:
        # Step 1: Get token and download existing document
        print("1. Downloading test.docx from SharePoint...")
        token = get_graph_token()
        file_stream = download_file(token, "/test.docx")
        doc = Document(file_stream)
        print("   ✅ Downloaded")

        # Step 2: Introduce violations
        print("\n2. Introducing style violations...")

        # Find or create Heading 1 paragraphs and mess them up
        violations_added = []

        for para in doc.paragraphs:
            if para.style.name == 'Heading 1':
                # Violation 1: Change font from Arial to Times New Roman
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                violations_added.append("Changed Heading 1 font to Times New Roman (should be Arial)")

                # Violation 2: Change color from RGB(0,51,153) to red
                for run in para.runs:
                    run.font.color.rgb = RGBColor(255, 0, 0)
                violations_added.append("Changed Heading 1 color to red (should be RGB(0,51,153))")

                break

        # If no Heading 1 exists, add one with violations
        if not violations_added:
            print("   No Heading 1 found, adding one with violations...")
            new_para = doc.add_paragraph('Test Heading with Violations', style='Heading 1')
            for run in new_para.runs:
                run.font.name = 'Times New Roman'
                run.font.color.rgb = RGBColor(255, 0, 0)
            violations_added.append("Added Heading 1 with Times New Roman font (should be Arial)")
            violations_added.append("Added Heading 1 with red color (should be RGB(0,51,153))")

        for violation in violations_added:
            print(f"   ⚠️  {violation}")

        # Step 3: Save modified document
        print("\n3. Saving modified document locally...")
        doc.save('test_with_violations.docx')
        print("   ✅ Saved to: test_with_violations.docx")

        # Step 4: Upload to SharePoint
        print("\n4. Uploading to SharePoint...")
        modified_stream = BytesIO()
        doc.save(modified_stream)
        modified_stream.seek(0)

        web_url = upload_file(token, modified_stream, "/test.docx")
        print(f"   ✅ Uploaded (overwrote original)")
        print(f"   URL: {web_url}")

        print("\n✨ Test document with violations is ready!")
        print(f"\nViolations introduced: {len(violations_added)}")
        print("\nNext: Run the validation to see auto-fix in action")

    except Exception as e:
        print(f"\n❌ Error: {str(e)}")
        import traceback
        traceback.print_exc()
        return 1

    return 0

if __name__ == "__main__":
    sys.exit(main())

"""
Create a new test document to upload and trigger Power Automate flow
"""
from docx import Document
from docx.shared import Pt, RGBColor

def main():
    print("📄 Creating new test document for Power Automate flow\n")

    # Create a new document
    doc = Document()

    # Add title
    doc.add_heading('Power Automate Test Document', 0)

    # Add a Heading 1 with VIOLATIONS
    heading = doc.add_paragraph('This Heading Has Style Violations', style='Heading 1')
    for run in heading.runs:
        run.font.name = 'Comic Sans MS'  # Wrong font
        run.font.color.rgb = RGBColor(0, 255, 0)  # Green instead of blue

    # Add some body text
    doc.add_paragraph('This is body text that should use Calibri font.')

    # Add another Heading 1 with violations
    heading2 = doc.add_paragraph('Another Heading with Problems', style='Heading 1')
    for run in heading2.runs:
        run.font.name = 'Times New Roman'  # Wrong font
        run.font.color.rgb = RGBColor(255, 0, 255)  # Purple instead of blue

    # Save locally
    filename = 'flow_test_document.docx'
    doc.save(filename)

    print(f"✅ Created: {filename}")
    print("\nViolations introduced:")
    print("  - Heading 1 uses Comic Sans MS (should be Arial)")
    print("  - Heading 1 uses green color (should be RGB(0,51,153))")
    print("  - Heading 1 uses Times New Roman (should be Arial)")
    print("  - Heading 1 uses purple color (should be RGB(0,51,153))")
    print("\n📤 Upload this file to SharePoint to trigger the Power Automate flow!")

if __name__ == "__main__":
    main()

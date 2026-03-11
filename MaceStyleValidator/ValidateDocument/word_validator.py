"""Word document (.docx) validation"""
import logging
from docx import Document
from docx.shared import RGBColor
from .ai_client import call_claude
from .enhanced_validators import validate_language_rules, validate_punctuation_rules, validate_grammar_rules


def _normalise_issue(item, rule=None):
    """Ensure an issue item is a structured dict"""
    if isinstance(item, dict) and 'rule_name' in item:
        return item
    return {
        'rule_name': rule.get('title', 'Unknown') if rule else 'Unknown',
        'rule_type': rule.get('rule_type', 'Unknown') if rule else 'Unknown',
        'description': str(item),
        'location': 'Document-wide',
        'priority': rule.get('priority', 999) if rule else 999
    }


def _normalise_fix(item, rule=None, changes=None):
    """Ensure a fix item is a structured dict"""
    if isinstance(item, dict) and 'rule_name' in item:
        if changes:
            item['changes'] = changes
        return item
    result = {
        'rule_name': rule.get('title', 'Unknown') if rule else 'Unknown',
        'rule_type': rule.get('rule_type', 'Unknown') if rule else 'Unknown',
        'found_value': 'Non-compliant value',
        'fixed_value': str(item),
        'location': 'Document-wide'
    }
    if changes:
        result['changes'] = changes
    return result


def validate_word_document(file_stream, rules):
    """Validate Word document against rules"""
    logging.info("Loading Word document...")
    doc = Document(file_stream)
    logging.info(f"Document loaded. Paragraphs: {len(doc.paragraphs)}, Tables: {len(doc.tables)}")

    issues = []
    fixes_applied = []

    word_rules = [r for r in rules if r['doc_type'] in ['Word', 'Both', 'All']]
    ai_rules = [r for r in word_rules if r.get('use_ai', False)]
    hard_coded_rules = [r for r in word_rules if not r.get('use_ai', False)]

    logging.info(f"AI rules: {len(ai_rules)}, Hard-coded rules: {len(hard_coded_rules)}")

    # AI-powered style corrections
    if ai_rules:
        try:
            full_text = "\n\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            if full_text.strip():
                result = call_claude(ai_rules, full_text)
                if result and result['changes_made'] > 0 and result['corrected_text']:
                    corrected_paras = result['corrected_text'].split('\n\n')
                    para_index = 0
                    ai_changes = []

                    for para in doc.paragraphs:
                        if para.text.strip() and para_index < len(corrected_paras):
                            original_text = para.text
                            if len(para.runs) > 0:
                                para.runs[0].text = corrected_paras[para_index]
                                for run in para.runs[1:]:
                                    run.text = ""
                            if original_text != corrected_paras[para_index]:
                                ai_changes.append({'before': original_text, 'after': corrected_paras[para_index], 'location': f'Paragraph {para_index + 1}'})
                            para_index += 1

                    ai_fix = {
                        'rule_name': 'AI Style Corrections',
                        'rule_type': 'AI',
                        'found_value': f'{result["changes_made"]} style violations',
                        'fixed_value': 'British English, contractions, symbols corrected',
                        'location': 'Document-wide'
                    }
                    if ai_changes:
                        ai_fix['changes'] = ai_changes
                    fixes_applied.append(ai_fix)
                    logging.info(f"Claude corrections applied: {result['changes_made']}")
        except Exception as e:
            logging.error(f"Claude validation failed: {e}")
            issues.append({
                'rule_name': 'AI Style Validation',
                'rule_type': 'AI',
                'description': f"AI validation failed: {e}",
                'location': 'N/A',
                'priority': 1
            })

    # Hard-coded rules
    for rule in hard_coded_rules:
        result = None
        if rule['rule_type'] == 'Font':
            result = _check_fonts(doc, rule)
        elif rule['rule_type'] == 'Color':
            result = _check_colors(doc, rule)
        elif rule['rule_type'] == 'Language':
            result = validate_language_rules(doc, rule)
        elif rule['rule_type'] == 'Grammar':
            result = validate_grammar_rules(doc, rule)
        elif rule['rule_type'] == 'Punctuation':
            result = validate_punctuation_rules(doc, rule)

        if result:
            for item in result.get('issues', []):
                issues.append(_normalise_issue(item, rule))
            result_changes = result.get('changes', [])
            for item in result.get('fixes', []):
                fixes_applied.append(_normalise_fix(item, rule, changes=result_changes))

    logging.info(f"Word validation complete. Issues: {len(issues)}, Fixes: {len(fixes_applied)}")
    return {'document': doc, 'issues': issues, 'fixes_applied': fixes_applied}


def _check_fonts(doc, rule):
    """Check and fix font issues in Word doc"""
    issues = []
    fixes = []
    expected_font = rule['expected_value']

    if rule['check_value'] == 'AllTextFont':
        issue_count = 0
        fix_count = 0

        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if run.text.strip():
                    if run.font.name is None or run.font.name != expected_font:
                        issue_count += 1
                        if rule['auto_fix']:
                            run.font.name = expected_font
                            fix_count += 1

        if issue_count > 0 and not rule['auto_fix']:
            issues.append({
                'rule_name': rule.get('title', 'All Text Font'),
                'rule_type': rule['rule_type'],
                'description': f"Found {issue_count} text runs with incorrect font (not {expected_font})",
                'location': 'Document-wide',
                'priority': rule.get('priority', 999)
            })
        if fix_count > 0:
            fixes.append({
                'rule_name': rule.get('title', 'All Text Font'),
                'rule_type': rule['rule_type'],
                'found_value': f'{issue_count} runs with wrong font',
                'fixed_value': expected_font,
                'location': f'Document-wide ({fix_count} runs)'
            })

    elif rule['check_value'] == 'Heading1Font':
        for para_idx, paragraph in enumerate(doc.paragraphs):
            if paragraph.style.name == 'Heading 1':
                current_font = paragraph.runs[0].font.name if paragraph.runs else None
                if current_font is None or current_font != expected_font:
                    if rule['auto_fix']:
                        for run in paragraph.runs:
                            run.font.name = expected_font
                        fixes.append({
                            'rule_name': rule.get('title', 'Heading 1 Font'),
                            'rule_type': rule['rule_type'],
                            'found_value': str(current_font),
                            'fixed_value': expected_font,
                            'location': f'Paragraph {para_idx + 1} (Heading 1)'
                        })
                    else:
                        issues.append({
                            'rule_name': rule.get('title', 'Heading 1 Font'),
                            'rule_type': rule['rule_type'],
                            'description': f"Heading 1 has incorrect font: {current_font}",
                            'location': f'Paragraph {para_idx + 1}',
                            'priority': rule.get('priority', 999)
                        })

    return {'issues': issues, 'fixes': fixes}


def _check_colors(doc, rule):
    """Check and fix color issues in Word doc"""
    issues = []
    fixes = []

    if rule['check_value'] == 'Heading1Color':
        expected_rgb = tuple(map(int, rule['expected_value'].split(',')))
        for para_idx, paragraph in enumerate(doc.paragraphs):
            if paragraph.style.name == 'Heading 1':
                for run in paragraph.runs:
                    if run.font.color.rgb:
                        current_rgb = run.font.color.rgb
                        if current_rgb != expected_rgb:
                            if rule['auto_fix']:
                                run.font.color.rgb = RGBColor(*expected_rgb)
                                fixes.append({
                                    'rule_name': rule.get('title', 'Heading 1 Color'),
                                    'rule_type': rule['rule_type'],
                                    'found_value': str(current_rgb),
                                    'fixed_value': str(expected_rgb),
                                    'location': f'Paragraph {para_idx + 1} (Heading 1)'
                                })
                            else:
                                issues.append({
                                    'rule_name': rule.get('title', 'Heading 1 Color'),
                                    'rule_type': rule['rule_type'],
                                    'description': f"Heading 1 color incorrect: {current_rgb}",
                                    'location': f'Paragraph {para_idx + 1}',
                                    'priority': rule.get('priority', 999)
                                })

    return {'issues': issues, 'fixes': fixes}

import azure.functions as func
import logging
import json
import os
from io import BytesIO
from datetime import datetime, timezone
from docx import Document
from docx.shared import Pt, RGBColor
from vsdx import VisioFile
import msal
import requests
from anthropic import Anthropic
from .enhanced_validators import validate_language_rules, validate_punctuation_rules, validate_grammar_rules
from .claude_validator import validate_with_claude
from .sharepoint_results import save_validation_result, update_document_metadata

# ============================================
# CONFIGURATION
# ============================================
def get_graph_token():
    """Get Microsoft Graph API access token using MSAL"""
    tenant_id = os.environ.get("SHAREPOINT_TENANT_ID", "2ab0866e-23d6-4688-be97-ce9f447135d8")
    client_id = os.environ.get("SHAREPOINT_CLIENT_ID", "c7859dae-6997-448f-9530-7166fe857e75")
    client_secret = os.environ.get("SHAREPOINT_CLIENT_SECRET", "DlD8Q~_NNgnpnVxKWsZTiz53DuNYrfrAjqkCDaP1")

    authority = f"https://login.microsoftonline.com/{tenant_id}"
    scope = ["https://graph.microsoft.com/.default"]

    app = msal.ConfidentialClientApplication(
        client_id,
        authority=authority,
        client_credential=client_secret
    )

    result = app.acquire_token_for_client(scopes=scope)

    if "access_token" in result:
        return result["access_token"]
    else:
        raise Exception(f"Failed to acquire token: {result.get('error_description', result)}")

def get_site_info():
    """Get SharePoint site information"""
    site_url = os.environ.get("SHAREPOINT_SITE_URL", "https://0rxf2.sharepoint.com/sites/StyleValidation")

    # Extract hostname and site path
    # Format: https://tenant.sharepoint.com/sites/sitename
    parts = site_url.replace("https://", "").split("/")
    hostname = parts[0]
    site_path = "/" + "/".join(parts[1:]) if len(parts) > 1 else ""

    return {
        "hostname": hostname,
        "site_path": site_path,
        "full_url": site_url
    }

def get_site_id(token):
    """Get SharePoint site ID using Graph API"""
    site_info = get_site_info()
    headers = {
        "Authorization": f"Bearer {token}",
        "Accept": "application/json"
    }

    site_url = f"https://graph.microsoft.com/v1.0/sites/{site_info['hostname']}:{site_info['site_path']}"
    site_response = requests.get(site_url, headers=headers)
    site_response.raise_for_status()
    return site_response.json()["id"]

# ============================================
# SHAREPOINT OPERATIONS
# ============================================
def fetch_validation_rules(token):
    """Fetch rules from SharePoint 'Style Rules' list using Graph API"""
    site_info = get_site_info()
    headers = {
        "Authorization": f"Bearer {token}",
        "Accept": "application/json"
    }

    # Get site ID first
    site_url = f"https://graph.microsoft.com/v1.0/sites/{site_info['hostname']}:{site_info['site_path']}"
    site_response = requests.get(site_url, headers=headers)
    site_response.raise_for_status()
    site_id = site_response.json()["id"]

    # Get list items
    list_url = f"https://graph.microsoft.com/v1.0/sites/{site_id}/lists/Style Rules/items?expand=fields"
    list_response = requests.get(list_url, headers=headers)
    list_response.raise_for_status()

    rules = []
    for item in list_response.json().get("value", []):
        fields = item.get("fields", {})
        rules.append({
            'title': fields.get('Title'),
            'rule_type': fields.get('RuleType'),
            'doc_type': fields.get('DocumentType'),
            'check_value': fields.get('CheckValue'),
            'expected_value': fields.get('ExpectedValue'),
            'auto_fix': fields.get('AutoFix'),
            'use_ai': fields.get('UseAI', False),  # Add UseAI field
            'priority': fields.get('Priority', 999)
        })

    rules.sort(key=lambda x: x['priority'])
    return rules

def build_dynamic_claude_prompt(ai_rules, document_text):
    """Build Claude prompt dynamically from SharePoint rules where UseAI=True"""

    # Group rules by type for better organization
    rules_by_type = {}
    for rule in ai_rules:
        rule_type = rule.get('rule_type', 'Other')
        if rule_type not in rules_by_type:
            rules_by_type[rule_type] = []
        rules_by_type[rule_type].append(rule)

    # Build rules description
    rules_description = []
    for rule_type, rules in sorted(rules_by_type.items()):
        rules_description.append(f"\n**{rule_type} Rules:**")
        for rule in rules:
            title = rule.get('title', 'Unknown rule')
            expected = rule.get('expected_value', '')
            if expected:
                rules_description.append(f"- {title} (use: {expected})")
            else:
                rules_description.append(f"- {title}")

    prompt = f"""You are a professional document editor applying the Mace Control Centre Writing Style Guide.

Apply ALL of the following corrections to the text:
{''.join(rules_description)}

Return a JSON object with two fields:
1. "corrected_text": the full corrected text (preserve paragraph breaks as \\n\\n)
2. "changes_made": total count of ALL changes made

Text to correct:
{document_text}"""

    return prompt

def download_file(token, file_path):
    """Download file from SharePoint using Graph API"""
    if not file_path:
        raise ValueError("file_path cannot be None or empty")

    logging.info(f"Downloading file: {file_path}")

    site_info = get_site_info()
    headers = {
        "Authorization": f"Bearer {token}",
        "Accept": "application/json"
    }

    # Get site ID
    site_url = f"https://graph.microsoft.com/v1.0/sites/{site_info['hostname']}:{site_info['site_path']}"
    site_response = requests.get(site_url, headers=headers)
    site_response.raise_for_status()
    site_id = site_response.json()["id"]

    # Convert server-relative URL to drive-relative path
    # Remove site path and "Shared Documents" from the path
    # E.g., "/sites/StyleValidation/Shared Documents/test.docx" -> "/test.docx"
    drive_relative_path = file_path
    if "Shared Documents/" in file_path:
        drive_relative_path = "/" + file_path.split("Shared Documents/", 1)[1]
    elif "Shared Documents" in file_path and file_path.endswith("Shared Documents"):
        # Handle case where path ends with "Shared Documents" (folder itself)
        drive_relative_path = "/"

    logging.info(f"Using drive-relative path: {drive_relative_path}")

    # Download file
    file_url = f"https://graph.microsoft.com/v1.0/sites/{site_id}/drive/root:{drive_relative_path}:/content"
    file_response = requests.get(file_url, headers=headers)
    file_response.raise_for_status()

    logging.info(f"File downloaded successfully, size: {len(file_response.content)} bytes")
    return BytesIO(file_response.content)

def upload_file(token, file_stream, target_path):
    """Upload file to SharePoint using Graph API"""
    if not target_path:
        raise ValueError("target_path cannot be None or empty")

    logging.info(f"Uploading file to: {target_path}")

    site_info = get_site_info()
    headers = {
        "Authorization": f"Bearer {token}",
        "Content-Type": "application/octet-stream"
    }

    # Get site ID
    site_url = f"https://graph.microsoft.com/v1.0/sites/{site_info['hostname']}:{site_info['site_path']}"
    site_response = requests.get(site_url, headers=headers.copy())
    site_response.raise_for_status()
    site_id = site_response.json()["id"]

    # Convert server-relative URL to drive-relative path
    # Remove site path and "Shared Documents" from the path
    # E.g., "/sites/StyleValidation/Shared Documents/test.docx" -> "/test.docx"
    drive_relative_path = target_path
    if "Shared Documents/" in target_path:
        drive_relative_path = "/" + target_path.split("Shared Documents/", 1)[1]
    elif "Shared Documents" in target_path and target_path.endswith("Shared Documents"):
        # Handle case where path ends with "Shared Documents" (folder itself)
        drive_relative_path = "/"

    logging.info(f"Using drive-relative path for upload: {drive_relative_path}")

    # Upload file
    file_stream.seek(0)
    upload_url = f"https://graph.microsoft.com/v1.0/sites/{site_id}/drive/root:{drive_relative_path}:/content"
    upload_response = requests.put(upload_url, headers=headers, data=file_stream.read())
    upload_response.raise_for_status()

    web_url = upload_response.json().get("webUrl")
    logging.info(f"File uploaded successfully: {web_url}")
    return web_url

def update_validation_status(token, item_id, status, report_url, list_name="Documents"):
    """Update ValidationStatus column in document library using Graph API"""
    site_info = get_site_info()
    headers = {
        "Authorization": f"Bearer {token}",
        "Content-Type": "application/json"
    }

    # Get site ID
    site_url = f"https://graph.microsoft.com/v1.0/sites/{site_info['hostname']}:{site_info['site_path']}"
    site_response = requests.get(site_url, headers=headers)
    site_response.raise_for_status()
    site_id = site_response.json()["id"]

    # Update list item
    update_url = f"https://graph.microsoft.com/v1.0/sites/{site_id}/lists/{list_name}/items/{item_id}/fields"
    update_data = {
        "ValidationStatus": status,
        "LastValidated": datetime.now(timezone.utc).isoformat()
    }
    if report_url:
        update_data["ValidationReport"] = report_url

    update_response = requests.patch(update_url, headers=headers, json=update_data)
    update_response.raise_for_status()

# ============================================
# VALIDATION LOGIC - WORD
# ============================================
def _normalise_issue(item, rule=None):
    """Ensure an issue item is a structured dict"""
    if isinstance(item, dict) and 'rule_name' in item:
        return item
    return {
        'rule_name': rule.get('title', 'Unknown') if rule else 'Unknown',
        'rule_type': rule.get('rule_type', 'Unknown') if rule else 'Unknown',
        'description': str(item),
        'location': 'Document-wide',
        'priority': rule.get('priority', 999) if rule else 999
    }

def _normalise_fix(item, rule=None):
    """Ensure a fix item is a structured dict"""
    if isinstance(item, dict) and 'rule_name' in item:
        return item
    return {
        'rule_name': rule.get('title', 'Unknown') if rule else 'Unknown',
        'rule_type': rule.get('rule_type', 'Unknown') if rule else 'Unknown',
        'found_value': 'Non-compliant value',
        'fixed_value': str(item),
        'location': 'Document-wide'
    }

def validate_word_document(file_stream, rules):
    """Validate Word document against rules"""
    logging.info("Loading Word document...")
    doc = Document(file_stream)
    logging.info(f"Document loaded. Paragraphs: {len(doc.paragraphs)}, Tables: {len(doc.tables)}")

    issues = []
    fixes_applied = []

    # Filter rules for Word documents
    word_rules = [r for r in rules if r['doc_type'] == 'Word']
    logging.info(f"Applying {len(word_rules)} Word validation rules")

    # Split rules into AI-enabled and hard-coded
    ai_rules = [r for r in word_rules if r.get('use_ai', False)]
    hard_coded_rules = [r for r in word_rules if not r.get('use_ai', False)]

    logging.info(f"AI-enabled rules: {len(ai_rules)}, Hard-coded rules: {len(hard_coded_rules)}")
    logging.info(f"AI rule titles: {[r.get('title', 'Unknown') for r in ai_rules]}")
    logging.info(f"Hard-coded rule titles: {[r.get('title', 'Unknown') for r in hard_coded_rules]}")

    # Process AI-enabled rules first (single Claude API call for all)
    if ai_rules:
        logging.info(f"Calling Claude AI for {len(ai_rules)} rules...")
        try:
            ai_result = validate_with_claude(doc, ai_rules)
            for item in ai_result.get('issues', []):
                issues.append(_normalise_issue(item, {'title': 'AI Style Check', 'rule_type': 'AI', 'priority': 1}))
            for item in ai_result.get('fixes_applied', []):
                fixes_applied.append(_normalise_fix(item, {'title': 'AI Style Check', 'rule_type': 'AI'}))
            logging.info(f"Claude validation complete: {len(ai_result.get('issues', []))} issues, {len(ai_result.get('fixes_applied', []))} fixes")
        except Exception as e:
            logging.error(f"Claude validation failed: {str(e)}")
            issues.append({
                'rule_name': 'AI Validation',
                'rule_type': 'AI',
                'description': f"AI validation failed: {str(e)}",
                'location': 'N/A',
                'priority': 1
            })

    # Process hard-coded rules
    for idx, rule in enumerate(hard_coded_rules):
        logging.info(f"Processing rule {idx+1}/{len(hard_coded_rules)}: {rule.get('title', 'Unknown')}")

        result = None
        if rule['rule_type'] == 'Font':
            result = check_word_fonts(doc, rule)
        elif rule['rule_type'] == 'Color':
            result = check_word_colors(doc, rule)
        elif rule['rule_type'] == 'Language':
            result = validate_language_rules(doc, rule)
        elif rule['rule_type'] == 'Grammar':
            result = validate_grammar_rules(doc, rule)
        elif rule['rule_type'] == 'Punctuation':
            result = validate_punctuation_rules(doc, rule)
        else:
            logging.info(f"Rule type '{rule['rule_type']}' not yet implemented")

        if result:
            for item in result.get('issues', []):
                issues.append(_normalise_issue(item, rule))
            for item in result.get('fixes', []):
                fixes_applied.append(_normalise_fix(item, rule))

    logging.info(f"Validation complete. Issues: {len(issues)}, Fixes: {len(fixes_applied)}")

    # DIAGNOSTIC LOGGING v2.3
    logging.info("=" * 60)
    logging.info("DIAGNOSTIC v2.3: validate_word_document returning")
    logging.info(f"DIAGNOSTIC: issues count = {len(issues)}")
    logging.info(f"DIAGNOSTIC: fixes_applied count = {len(fixes_applied)}")
    if issues:
        logging.info(f"DIAGNOSTIC: issues[0] = {issues[0]}")
    if fixes_applied:
        logging.info(f"DIAGNOSTIC: fixes_applied[0] = {fixes_applied[0]}")
    logging.info("=" * 60)

    return {
        'document': doc,
        'issues': issues,
        'fixes_applied': fixes_applied
    }

def check_word_fonts(doc, rule):
    """Check and fix font issues in Word doc"""
    issues = []
    fixes = []
    expected_font = rule['expected_value']

    logging.info(f"Checking fonts: {rule['check_value']} -> {expected_font}")

    # Check all text fonts
    if rule['check_value'] == 'AllTextFont':
        issue_count = 0
        fix_count = 0

        for para_idx, paragraph in enumerate(doc.paragraphs):
            for run in paragraph.runs:
                if run.text.strip():  # Only check runs with actual text
                    current_font = run.font.name

                    # Handle None font names or mismatches
                    if current_font is None or current_font != expected_font:
                        issue_count += 1

                        if rule['auto_fix']:
                            run.font.name = expected_font
                            fix_count += 1

        if issue_count > 0 and not rule['auto_fix']:
            issues.append({
                'rule_name': rule.get('title', 'All Text Font'),
                'rule_type': rule['rule_type'],
                'description': f"Found {issue_count} text runs with incorrect font (not {expected_font})",
                'location': 'Document-wide',
                'priority': rule.get('priority', 999)
            })
        if fix_count > 0:
            fixes.append({
                'rule_name': rule.get('title', 'All Text Font'),
                'rule_type': rule['rule_type'],
                'found_value': f'{issue_count} runs with wrong font',
                'fixed_value': expected_font,
                'location': f'Document-wide ({fix_count} runs)'
            })

        logging.info(f"Font check complete: {issue_count} issues, {fix_count} fixes")

    # Check Heading 1 font
    elif rule['check_value'] == 'Heading1Font':
        for para_idx, paragraph in enumerate(doc.paragraphs):
            if paragraph.style.name == 'Heading 1':
                current_font = paragraph.runs[0].font.name if paragraph.runs else None

                if current_font is None or current_font != expected_font:
                    if rule['auto_fix']:
                        for run in paragraph.runs:
                            run.font.name = expected_font
                        fixes.append({
                            'rule_name': rule.get('title', 'Heading 1 Font'),
                            'rule_type': rule['rule_type'],
                            'found_value': str(current_font),
                            'fixed_value': expected_font,
                            'location': f'Paragraph {para_idx + 1} (Heading 1)'
                        })
                    else:
                        issues.append({
                            'rule_name': rule.get('title', 'Heading 1 Font'),
                            'rule_type': rule['rule_type'],
                            'description': f"Heading 1 has incorrect font: {current_font}",
                            'location': f'Paragraph {para_idx + 1}',
                            'priority': rule.get('priority', 999)
                        })

    return {'issues': issues, 'fixes': fixes}

def check_word_colors(doc, rule):
    """Check and fix color issues in Word doc"""
    issues = []
    fixes = []

    # Example: Check heading color
    if rule['check_value'] == 'Heading1Color':
        # Parse expected RGB from rule (e.g., "0,51,153")
        expected_rgb = tuple(map(int, rule['expected_value'].split(',')))

        for para_idx, paragraph in enumerate(doc.paragraphs):
            if paragraph.style.name == 'Heading 1':
                for run in paragraph.runs:
                    if run.font.color.rgb:
                        current_rgb = run.font.color.rgb

                        if current_rgb != expected_rgb:
                            if rule['auto_fix']:
                                run.font.color.rgb = RGBColor(*expected_rgb)
                                fixes.append({
                                    'rule_name': rule.get('title', 'Heading 1 Color'),
                                    'rule_type': rule['rule_type'],
                                    'found_value': str(current_rgb),
                                    'fixed_value': str(expected_rgb),
                                    'location': f'Paragraph {para_idx + 1} (Heading 1)'
                                })
                            else:
                                issues.append({
                                    'rule_name': rule.get('title', 'Heading 1 Color'),
                                    'rule_type': rule['rule_type'],
                                    'description': f"Heading 1 color incorrect: {current_rgb}",
                                    'location': f'Paragraph {para_idx + 1}',
                                    'priority': rule.get('priority', 999)
                                })

    return {'issues': issues, 'fixes': fixes}

# ============================================
# VALIDATION LOGIC - VISIO
# ============================================
def validate_visio_document(file_stream, rules):
    """Validate Visio document against rules"""
    logging.info("Loading Visio document...")
    visio = VisioFile(file_stream)

    page_count = len(visio.pages)
    logging.info(f"Visio document loaded. Pages: {page_count}")

    issues = []
    fixes_applied = []

    # Filter rules for Visio documents (including 'Both')
    visio_rules = [r for r in rules if r['doc_type'] in ['Visio', 'Both']]
    logging.info(f"Applying {len(visio_rules)} Visio validation rules")

    # Split rules into AI-enabled and hard-coded
    ai_rules = [r for r in visio_rules if r.get('use_ai', False)]
    hard_coded_rules = [r for r in visio_rules if not r.get('use_ai', False)]

    logging.info(f"AI-enabled rules: {len(ai_rules)}, Hard-coded rules: {len(hard_coded_rules)}")

    # Extract all text from Visio shapes
    shape_texts = []
    for page in visio.pages:
        shape_texts.extend(extract_shape_texts(page, page.child_shapes))

    logging.info(f"Extracted text from {len(shape_texts)} shapes")

    # Process AI-enabled rules if we have text
    if ai_rules and shape_texts:
        logging.info(f"Calling Claude AI for {len(ai_rules)} Visio rules...")
        try:
            # Combine all text for AI validation
            combined_text = "\n\n".join([st['text'] for st in shape_texts if st['text'].strip()])

            if combined_text.strip():
                # Get Anthropic API key
                api_key = os.environ.get("ANTHROPIC_API_KEY")
                if api_key:
                    from anthropic import Anthropic
                    client = Anthropic(api_key=api_key)

                    # Build dynamic prompt from rules
                    prompt = build_dynamic_claude_prompt(ai_rules, combined_text)

                    response = client.messages.create(
                        model="claude-3-haiku-20240307",
                        max_tokens=4096,
                        temperature=0.3,
                        messages=[{"role": "user", "content": prompt}]
                    )

                    response_text = response.content[0].text
                    json_start = response_text.find('{')
                    json_end = response_text.rfind('}') + 1

                    if json_start >= 0 and json_end > json_start:
                        result_json = json.loads(response_text[json_start:json_end])
                        corrected_text = result_json.get('corrected_text', '')
                        changes_count = result_json.get('changes_made', 0)

                        if changes_count > 0 and corrected_text:
                            # Split corrected text back into shape texts
                            corrected_parts = corrected_text.split('\n\n')

                            # Apply corrections to shapes
                            shapes_updated = 0
                            for idx, shape_data in enumerate(shape_texts):
                                if idx < len(corrected_parts):
                                    new_text = corrected_parts[idx].strip()
                                    if new_text != shape_data['text'].strip():
                                        shape_data['shape'].text = new_text
                                        shapes_updated += 1

                            issues.append(f"Found {changes_count} style violations in Visio shapes")
                            fixes_applied.append(f"Applied {changes_count} style corrections to {shapes_updated} shapes")
                            logging.info(f"Claude corrections applied: {changes_count} changes in {shapes_updated} shapes")
                        else:
                            logging.info("No AI corrections needed")
                else:
                    logging.warning("ANTHROPIC_API_KEY not set - skipping AI validation")

        except Exception as e:
            logging.error(f"Claude validation failed for Visio: {str(e)}")
            issues.append(f"AI validation failed: {str(e)}")

    # Process hard-coded rules
    for idx, rule in enumerate(hard_coded_rules):
        logging.info(f"Processing Visio rule {idx+1}/{len(hard_coded_rules)}: {rule.get('title', 'Unknown')}")

        if rule['rule_type'] == 'Color':
            result = check_visio_colors(visio, rule)
            issues.extend(result['issues'])
            fixes_applied.extend(result['fixes'])

        elif rule['rule_type'] == 'Font':
            result = check_visio_fonts(visio, rule)
            issues.extend(result['issues'])
            fixes_applied.extend(result['fixes'])

        elif rule['rule_type'] == 'Size':
            result = check_visio_shape_size(visio, rule)
            issues.extend(result['issues'])
            fixes_applied.extend(result['fixes'])

        elif rule['rule_type'] == 'Position':
            result = check_visio_position(visio, rule)
            issues.extend(result['issues'])
            fixes_applied.extend(result['fixes'])

        elif rule['rule_type'] == 'PageDimensions':
            result = check_visio_page_dimensions(visio, rule)
            issues.extend(result['issues'])
            fixes_applied.extend(result['fixes'])

    logging.info(f"Visio validation complete. Issues: {len(issues)}, Fixes: {len(fixes_applied)}")

    return {
        'document': visio,
        'issues': issues,
        'fixes_applied': fixes_applied
    }

def extract_shape_texts(page, shapes, shape_list=None):
    """Recursively extract text from all shapes in a page"""
    if shape_list is None:
        shape_list = []

    for shape in shapes:
        # Get shape text
        if hasattr(shape, 'text') and shape.text:
            text = str(shape.text).strip()
            if text:
                shape_list.append({
                    'shape': shape,
                    'text': text,
                    'page': page.name
                })

        # Recursively process child shapes
        if hasattr(shape, 'child_shapes') and shape.child_shapes:
            extract_shape_texts(page, shape.child_shapes, shape_list)

    return shape_list

def check_visio_colors(visio, rule):
    """Check and fix colors in Visio diagrams"""
    issues = []
    fixes = []
    expected_color = rule['expected_value']  # e.g., "#003399"

    logging.info(f"Checking Visio colors: {rule['check_value']} -> {expected_color}")

    issue_count = 0
    fix_count = 0

    def process_shapes_for_color(shapes, parent_name=""):
        """Recursively process shapes for color validation"""
        nonlocal issue_count, fix_count

        for shape in shapes:
            # Only process shapes with text (visible shapes)
            if hasattr(shape, 'text') and shape.text and str(shape.text).strip():
                try:
                    # Check fill color
                    if rule['check_value'] == 'ShapeFillColor':
                        current_color = getattr(shape, 'fill_color', None)
                        if current_color and current_color != expected_color:
                            issue_count += 1
                            if rule['auto_fix']:
                                shape.fill_color = expected_color
                                fix_count += 1

                    # Check text color
                    elif rule['check_value'] == 'ShapeTextColor':
                        current_color = getattr(shape, 'text_color', None)
                        if current_color and current_color != expected_color:
                            issue_count += 1
                            if rule['auto_fix']:
                                shape.text_color = expected_color
                                fix_count += 1

                except Exception as e:
                    logging.warning(f"Could not check/set color for shape: {str(e)}")

            # Recursively process child shapes
            if hasattr(shape, 'child_shapes') and shape.child_shapes:
                process_shapes_for_color(shape.child_shapes, f"{parent_name}/child")

    # Process all pages
    for page in visio.pages:
        if hasattr(page, 'child_shapes'):
            process_shapes_for_color(page.child_shapes, page.name)

    if issue_count > 0:
        issues.append(f"Found {issue_count} shapes with incorrect {rule['check_value']}")
    if fix_count > 0:
        fixes.append(f"Fixed {fix_count} shapes to {expected_color}")

    logging.info(f"Color check complete: {issue_count} issues, {fix_count} fixes")
    return {'issues': issues, 'fixes': fixes}

def check_visio_fonts(visio, rule):
    """Check and fix fonts in Visio diagrams"""
    issues = []
    fixes = []
    expected_font = rule['expected_value']

    logging.info(f"Checking Visio fonts: {rule['check_value']} -> {expected_font}")

    issue_count = 0
    fix_count = 0

    def process_shapes_for_font(shapes, parent_name=""):
        """Recursively process shapes for font validation"""
        nonlocal issue_count, fix_count

        for shape in shapes:
            # Only process shapes with text
            if hasattr(shape, 'text') and shape.text and str(shape.text).strip():
                try:
                    # Attempt to check font using cell values
                    # In Visio, font 0 = default font (Arial)
                    # Font references are stored in Character section cells

                    # Check if shape has font property or cell
                    if rule['check_value'] == 'AllTextFont':
                        # Try to get current font value from Character.Font cell
                        try:
                            # The vsdx library may allow access to cells
                            current_font = shape.cells.get('Char.Font', None)

                            # If we can read the font and it's not default (0 for Arial)
                            if current_font is not None:
                                font_value = current_font.value if hasattr(current_font, 'value') else str(current_font)

                                # Set font to 0 (Arial/default) if it's different
                                if font_value != '0':
                                    issue_count += 1
                                    if rule['auto_fix']:
                                        # Try to set font to 0 (Arial)
                                        shape.set_cell_value('Char.Font', '0')
                                        fix_count += 1
                        except (AttributeError, KeyError):
                            # If cells property doesn't exist or method not available,
                            # try alternative approach using set_cell_value directly
                            if rule['auto_fix']:
                                try:
                                    # Attempt to set font to default (0 = Arial)
                                    shape.set_cell_value('Char.Font', '0')
                                    fix_count += 1
                                    # Count as issue if we're setting it
                                    issue_count += 1
                                except Exception as set_error:
                                    logging.debug(f"Could not set font via set_cell_value: {str(set_error)}")

                except Exception as e:
                    logging.warning(f"Could not check/set font for shape: {str(e)}")

            # Recursively process child shapes
            if hasattr(shape, 'child_shapes') and shape.child_shapes:
                process_shapes_for_font(shape.child_shapes, f"{parent_name}/child")

    # Process all pages
    for page in visio.pages:
        if hasattr(page, 'child_shapes'):
            process_shapes_for_font(page.child_shapes, page.name)

    if issue_count > 0:
        issues.append(f"Found {issue_count} shapes with incorrect font")
    if fix_count > 0:
        fixes.append(f"Fixed {fix_count} shapes to {expected_font}")

    logging.info(f"Font check complete: {issue_count} issues, {fix_count} fixes")
    return {'issues': issues, 'fixes': fixes}

def check_visio_shape_size(visio, rule):
    """Check and fix shape dimensions in Visio diagrams"""
    issues = []
    fixes = []

    logging.info(f"Checking Visio shape sizes: {rule['check_value']}")

    # Parse expected value (format: "WIDTHxHEIGHT" e.g., "3.0x1.0")
    expected_value = rule['expected_value']
    tolerance = float(rule.get('tolerance', 0.1))  # Default ±0.1 inch tolerance

    try:
        if 'x' in expected_value.lower():
            expected_width, expected_height = map(float, expected_value.lower().split('x'))
        else:
            logging.warning(f"Invalid size format: {expected_value}. Expected format: WIDTHxHEIGHT")
            return {'issues': issues, 'fixes': fixes}
    except ValueError:
        logging.warning(f"Could not parse size value: {expected_value}")
        return {'issues': issues, 'fixes': fixes}

    issue_count = 0
    fix_count = 0

    def process_shapes_for_size(shapes, parent_name=""):
        """Recursively process shapes for size validation"""
        nonlocal issue_count, fix_count

        for shape in shapes:
            # Only process shapes with text (visible shapes)
            if hasattr(shape, 'text') and shape.text and str(shape.text).strip():
                try:
                    # Get current dimensions
                    current_width = getattr(shape, 'width', None)
                    current_height = getattr(shape, 'height', None)

                    if current_width is not None and current_height is not None:
                        # Check if dimensions are within tolerance
                        width_diff = abs(current_width - expected_width)
                        height_diff = abs(current_height - expected_height)

                        if width_diff > tolerance or height_diff > tolerance:
                            issue_count += 1

                            if rule['auto_fix']:
                                # Set new dimensions
                                shape.width = expected_width
                                shape.height = expected_height
                                fix_count += 1

                except Exception as e:
                    logging.warning(f"Could not check/set size for shape: {str(e)}")

            # Recursively process child shapes
            if hasattr(shape, 'child_shapes') and shape.child_shapes:
                process_shapes_for_size(shape.child_shapes, f"{parent_name}/child")

    # Process all pages
    for page in visio.pages:
        if hasattr(page, 'child_shapes'):
            process_shapes_for_size(page.child_shapes, page.name)

    if issue_count > 0:
        issues.append(f"Found {issue_count} shapes with incorrect dimensions (expected {expected_width}x{expected_height})")
    if fix_count > 0:
        fixes.append(f"Resized {fix_count} shapes to {expected_width}x{expected_height}")

    logging.info(f"Size check complete: {issue_count} issues, {fix_count} fixes")
    return {'issues': issues, 'fixes': fixes}

def check_visio_position(visio, rule):
    """Check and fix shape positions in Visio diagrams"""
    issues = []
    fixes = []

    logging.info(f"Checking Visio shape positions: {rule['check_value']}")

    # Parse rule parameters
    # CheckValue options: TopMargin, LeftMargin, BottomMargin, RightMargin, ExactPosition
    check_type = rule['check_value']
    expected_value = rule['expected_value']
    tolerance = float(rule.get('tolerance', 0.1))  # Default ±0.1 inch tolerance

    issue_count = 0
    fix_count = 0

    def process_shapes_for_position(shapes, parent_name=""):
        """Recursively process shapes for position validation"""
        nonlocal issue_count, fix_count

        for shape in shapes:
            # Only process shapes with text (visible shapes)
            if hasattr(shape, 'text') and shape.text and str(shape.text).strip():
                try:
                    current_x = getattr(shape, 'x', None)
                    current_y = getattr(shape, 'y', None)

                    if current_x is None or current_y is None:
                        continue

                    # Validate based on check type
                    if check_type == 'TopMargin':
                        # Check if shape is within top margin (max Y value)
                        max_y = float(expected_value)
                        if current_y > max_y + tolerance:
                            issue_count += 1
                            if rule['auto_fix']:
                                shape.y = max_y
                                fix_count += 1

                    elif check_type == 'LeftMargin':
                        # Check if shape is beyond left margin (min X value)
                        min_x = float(expected_value)
                        if current_x < min_x - tolerance:
                            issue_count += 1
                            if rule['auto_fix']:
                                shape.x = min_x
                                fix_count += 1

                    elif check_type == 'RightMargin':
                        # Check if shape is beyond right margin (max X value)
                        max_x = float(expected_value)
                        if current_x > max_x + tolerance:
                            issue_count += 1
                            if rule['auto_fix']:
                                shape.x = max_x
                                fix_count += 1

                    elif check_type == 'BottomMargin':
                        # Check if shape is below bottom margin (min Y value)
                        min_y = float(expected_value)
                        if current_y < min_y - tolerance:
                            issue_count += 1
                            if rule['auto_fix']:
                                shape.y = min_y
                                fix_count += 1

                    elif check_type == 'ExactPosition':
                        # Check exact X,Y position (format: "X,Y")
                        expected_x, expected_y = map(float, expected_value.split(','))
                        x_diff = abs(current_x - expected_x)
                        y_diff = abs(current_y - expected_y)

                        if x_diff > tolerance or y_diff > tolerance:
                            issue_count += 1
                            if rule['auto_fix']:
                                shape.x = expected_x
                                shape.y = expected_y
                                fix_count += 1

                except Exception as e:
                    logging.warning(f"Could not check/set position for shape: {str(e)}")

            # Recursively process child shapes
            if hasattr(shape, 'child_shapes') and shape.child_shapes:
                process_shapes_for_position(shape.child_shapes, f"{parent_name}/child")

    # Process all pages
    for page in visio.pages:
        if hasattr(page, 'child_shapes'):
            process_shapes_for_position(page.child_shapes, page.name)

    if issue_count > 0:
        issues.append(f"Found {issue_count} shapes with incorrect position ({check_type})")
    if fix_count > 0:
        fixes.append(f"Repositioned {fix_count} shapes for {check_type}")

    logging.info(f"Position check complete: {issue_count} issues, {fix_count} fixes")
    return {'issues': issues, 'fixes': fixes}

def check_visio_page_dimensions(visio, rule):
    """Check and fix page dimensions in Visio diagrams"""
    issues = []
    fixes = []

    logging.info(f"Checking Visio page dimensions: {rule['check_value']}")

    # Parse expected value (format: "WIDTHxHEIGHT" e.g., "11.0x8.5")
    expected_value = rule['expected_value']

    try:
        if 'x' in expected_value.lower():
            expected_width, expected_height = map(float, expected_value.lower().split('x'))
        else:
            logging.warning(f"Invalid page size format: {expected_value}. Expected format: WIDTHxHEIGHT")
            return {'issues': issues, 'fixes': fixes}
    except ValueError:
        logging.warning(f"Could not parse page size value: {expected_value}")
        return {'issues': issues, 'fixes': fixes}

    issue_count = 0
    fix_count = 0

    # Check all pages
    for page in visio.pages:
        try:
            current_width = getattr(page, 'width', None)
            current_height = getattr(page, 'height', None)

            if current_width is not None and current_height is not None:
                # Check if dimensions match
                if current_width != expected_width or current_height != expected_height:
                    issue_count += 1

                    if rule['auto_fix']:
                        # Set new page dimensions
                        page.width = expected_width
                        page.height = expected_height
                        fix_count += 1

        except Exception as e:
            logging.warning(f"Could not check/set page dimensions for page '{page.name}': {str(e)}")

    if issue_count > 0:
        issues.append(f"Found {issue_count} pages with incorrect dimensions (expected {expected_width}x{expected_height})")
    if fix_count > 0:
        fixes.append(f"Resized {fix_count} pages to {expected_width}x{expected_height}")

    logging.info(f"Page dimensions check complete: {issue_count} issues, {fix_count} fixes")
    return {'issues': issues, 'fixes': fixes}

# ============================================
# REPORT GENERATION
# ============================================
def _escape_html(text):
    """Escape HTML special characters"""
    if text is None:
        return ''
    return str(text).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;').replace('"', '&quot;')

def generate_report(file_name, issues, fixes_applied):
    """Generate validation report as HTML with Mace branding.

    Args:
        file_name: Name of the validated document.
        issues: List of dicts with keys: rule_name, rule_type, description, location, priority.
        fixes_applied: List of dicts with keys: rule_name, rule_type, found_value, fixed_value, location.
    """
    # Separate remaining issues (items in issues that were NOT fixed)
    remaining_issues = [i for i in issues if isinstance(i, dict)]
    total_issues_found = len(remaining_issues) + len(fixes_applied)
    status = "Passed" if len(remaining_issues) == 0 else "Failed"
    status_color = "#28a745" if status == "Passed" else "#dc3545"
    validation_time = datetime.now(timezone.utc).strftime('%d %B %Y at %H:%M:%S UTC')

    # Build fixes table rows
    fixes_rows = ''
    for fix in fixes_applied:
        if isinstance(fix, dict):
            fixes_rows += f"""<tr>
                <td>{_escape_html(fix.get('rule_name', ''))}</td>
                <td><span class="rule-type-badge">{_escape_html(fix.get('rule_type', ''))}</span></td>
                <td>{_escape_html(fix.get('found_value', ''))}</td>
                <td>{_escape_html(fix.get('fixed_value', ''))}</td>
                <td>{_escape_html(fix.get('location', ''))}</td>
            </tr>"""
        else:
            fixes_rows += f"""<tr>
                <td colspan="5">{_escape_html(str(fix))}</td>
            </tr>"""

    # Build remaining issues table rows
    issues_rows = ''
    for issue in remaining_issues:
        priority = issue.get('priority', 999)
        priority_label = 'High' if priority <= 3 else ('Medium' if priority <= 6 else 'Low')
        priority_color = '#dc3545' if priority <= 3 else ('#f0ad4e' if priority <= 6 else '#6c757d')
        issues_rows += f"""<tr>
            <td>{_escape_html(issue.get('rule_name', ''))}</td>
            <td><span class="rule-type-badge">{_escape_html(issue.get('rule_type', ''))}</span></td>
            <td>{_escape_html(issue.get('description', ''))}</td>
            <td>{_escape_html(issue.get('location', ''))}</td>
            <td><span style="color:{priority_color};font-weight:bold;">{priority_label}</span></td>
        </tr>"""

    # Build the fixes section
    if fixes_applied:
        fixes_section = f"""<div class="section">
            <h2>Changes Made ({len(fixes_applied)})</h2>
            <table>
                <thead><tr>
                    <th>Rule Name</th><th>Rule Type</th><th>What Was Found</th><th>What It Was Changed To</th><th>Location</th>
                </tr></thead>
                <tbody>{fixes_rows}</tbody>
            </table>
        </div>"""
    else:
        fixes_section = ''

    # Build the remaining issues section
    if remaining_issues:
        issues_section = f"""<div class="section">
            <h2>Remaining Issues ({len(remaining_issues)})</h2>
            <table>
                <thead><tr>
                    <th>Rule Name</th><th>Rule Type</th><th>Issue Description</th><th>Location</th><th>Priority</th>
                </tr></thead>
                <tbody>{issues_rows}</tbody>
            </table>
        </div>"""
    else:
        issues_section = ''

    # No-issues confirmation
    if total_issues_found == 0:
        no_issues_section = """<div class="section passed-section">
            <h2>All Clear</h2>
            <p>This document fully complies with the Mace Control Centre Writing Style Guide. No issues were found.</p>
        </div>"""
    else:
        no_issues_section = ''

    report_html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Validation Report - {_escape_html(file_name)}</title>
    <style>
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        body {{
            font-family: Arial, Helvetica, sans-serif;
            max-width: 1100px;
            margin: 0 auto;
            padding: 24px;
            background: #f4f6f9;
            color: #333;
            line-height: 1.5;
        }}
        .header {{
            background: linear-gradient(135deg, #1F4E79 0%, #1671CB 100%);
            color: #fff;
            padding: 32px;
            border-radius: 8px;
            margin-bottom: 24px;
        }}
        .header h1 {{
            font-size: 22px;
            margin-bottom: 16px;
            font-weight: 700;
        }}
        .status-badge {{
            display: inline-block;
            padding: 6px 18px;
            background: {status_color};
            color: #fff;
            border-radius: 16px;
            font-weight: 700;
            font-size: 13px;
            text-transform: uppercase;
            letter-spacing: 0.5px;
        }}
        .meta-info {{
            margin-top: 14px;
            font-size: 13px;
            opacity: 0.92;
        }}
        .meta-info div {{ margin-bottom: 3px; }}
        .summary {{
            background: #fff;
            padding: 24px;
            border-radius: 8px;
            margin-bottom: 20px;
            box-shadow: 0 1px 3px rgba(0,0,0,0.08);
        }}
        .summary h2 {{
            font-size: 16px;
            color: #1F4E79;
            margin-bottom: 16px;
            padding-bottom: 8px;
            border-bottom: 2px solid #1F4E79;
        }}
        .summary-grid {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(180px, 1fr));
            gap: 16px;
        }}
        .summary-card {{
            background: #f4f6f9;
            padding: 18px;
            border-radius: 6px;
            text-align: center;
            border-left: 4px solid #1F4E79;
        }}
        .summary-card .number {{
            font-size: 32px;
            font-weight: 700;
            color: #1F4E79;
            margin-bottom: 4px;
        }}
        .summary-card .label {{
            font-size: 12px;
            color: #6c757d;
            text-transform: uppercase;
            letter-spacing: 0.5px;
        }}
        .section {{
            background: #fff;
            padding: 24px;
            border-radius: 8px;
            margin-bottom: 20px;
            box-shadow: 0 1px 3px rgba(0,0,0,0.08);
        }}
        .section h2 {{
            font-size: 16px;
            color: #1F4E79;
            margin-top: 0;
            margin-bottom: 16px;
            padding-bottom: 8px;
            border-bottom: 2px solid #1F4E79;
        }}
        .passed-section {{
            border-left: 4px solid #28a745;
        }}
        .passed-section h2 {{
            color: #28a745;
            border-bottom-color: #28a745;
        }}
        .passed-section p {{
            color: #555;
            padding: 20px 0;
            text-align: center;
            font-size: 15px;
        }}
        table {{
            width: 100%;
            border-collapse: collapse;
            font-size: 13px;
        }}
        thead th {{
            background: #1F4E79;
            color: #fff;
            padding: 10px 12px;
            text-align: left;
            font-weight: 600;
            font-size: 12px;
            text-transform: uppercase;
            letter-spacing: 0.3px;
        }}
        thead th:first-child {{ border-radius: 4px 0 0 0; }}
        thead th:last-child {{ border-radius: 0 4px 0 0; }}
        tbody td {{
            padding: 10px 12px;
            border-bottom: 1px solid #e9ecef;
            vertical-align: top;
        }}
        tbody tr:hover {{ background: #f8f9fb; }}
        tbody tr:last-child td {{ border-bottom: none; }}
        .rule-type-badge {{
            display: inline-block;
            padding: 2px 8px;
            background: #e8f0fe;
            color: #1671CB;
            border-radius: 10px;
            font-size: 11px;
            font-weight: 600;
        }}
        .footer {{
            text-align: center;
            margin-top: 28px;
            padding: 16px;
            color: #999;
            font-size: 11px;
        }}
        .footer span {{ color: #1F4E79; font-weight: 600; }}
    </style>
</head>
<body>
    <div class="header">
        <h1>Mace Style Validation Report</h1>
        <span class="status-badge">{status}</span>
        <div class="meta-info">
            <div><strong>Document:</strong> {_escape_html(file_name)}</div>
            <div><strong>Validated:</strong> {validation_time}</div>
        </div>
    </div>

    <div class="summary">
        <h2>Summary</h2>
        <div class="summary-grid">
            <div class="summary-card">
                <div class="number">{total_issues_found}</div>
                <div class="label">Issues Found</div>
            </div>
            <div class="summary-card">
                <div class="number">{len(fixes_applied)}</div>
                <div class="label">Auto-Fixed</div>
            </div>
            <div class="summary-card">
                <div class="number">{len(remaining_issues)}</div>
                <div class="label">Remaining</div>
            </div>
        </div>
    </div>

    {no_issues_section}
    {fixes_section}
    {issues_section}

    <div class="footer">
        <span>Mace Style Validator</span> &middot; Control Centre Writing Style Guide<br>
        Powered by Azure Functions &amp; Claude AI
    </div>
</body>
</html>"""
    return report_html

# ============================================
# MAIN FUNCTION
# ============================================
def main(req: func.HttpRequest) -> func.HttpResponse:
    logging.info('=== STYLE VALIDATION FUNCTION TRIGGERED ===')

    try:
        # 1. Parse request
        logging.info('Step 1: Parsing request...')
        req_body = req.get_json()
        logging.info(f"Request keys: {list(req_body.keys())}")

        item_id = req_body.get('itemId') or req_body.get('ID')

        # Try multiple possible parameter names for file name
        file_name = (req_body.get('fileName') or
                    req_body.get('FileLeafRef') or
                    req_body.get('Name'))

        file_extension = os.path.splitext(file_name)[1].lower() if file_name else ''

        # Check if file content is provided directly (base64 encoded)
        file_content_base64 = req_body.get('fileContent')

        # Try multiple possible parameter names for file URL (for legacy support)
        file_url = (req_body.get('fileUrl') or
                   req_body.get('FileRef') or
                   req_body.get('ServerRelativeUrl') or
                   req_body.get('fileRef'))

        logging.info(f"Request params - File: {file_name}, ID: {item_id}, Has content: {bool(file_content_base64)}, URL: {file_url}")

        # 2. Get Microsoft Graph API token
        logging.info('Step 2: Acquiring Graph API token...')
        token = get_graph_token()
        logging.info('Token acquired successfully')

        # 3. Update status to "Validating..."
        logging.info('Step 3: Updating validation status...')
        update_validation_status(token, item_id, "Validating...", None)

        # 4. Fetch validation rules
        logging.info('Step 4: Fetching validation rules...')
        rules = fetch_validation_rules(token)
        logging.info(f"Loaded {len(rules)} validation rules")

        # 5. Get file content
        if file_content_base64:
            # File content provided directly as base64
            logging.info('Step 5: Decoding file content from base64...')
            import base64
            file_bytes = base64.b64decode(file_content_base64)
            file_stream = BytesIO(file_bytes)
            logging.info(f"File decoded successfully, size: {len(file_bytes)} bytes")
        elif file_url:
            # Download file from SharePoint using Graph API
            logging.info('Step 5: Downloading file from SharePoint...')
            file_stream = download_file(token, file_url)
        else:
            raise ValueError("Either fileContent or fileUrl must be provided")
        
        # 6. Validate based on file type
        logging.info(f'Step 6: Validating document type {file_extension}...')
        if file_extension in ['.docx', '.doc']:
            logging.info("=" * 60)
            logging.info("v4.2: Dynamic Rules + Validation Results + Document Library Linkback")
            logging.info("=" * 60)

            from docx import Document
            doc = Document(file_stream)

            issues = []
            fixes_applied = []

            # 1. Use Claude for comprehensive style corrections
            api_key = os.environ.get("ANTHROPIC_API_KEY")
            if api_key:
                logging.info("Calling Claude for style corrections...")

                # Extract all text from document
                full_text = "\n\n".join([p.text for p in doc.paragraphs if p.text.strip()])

                if full_text.strip():
                    try:
                        # Fetch rules from SharePoint and filter for AI-enabled rules
                        logging.info("Fetching validation rules from SharePoint...")
                        all_rules = fetch_validation_rules(token)
                        ai_rules = [r for r in all_rules if r.get('use_ai', False)]
                        logging.info(f"Found {len(ai_rules)} AI-enabled rules from SharePoint")

                        client = Anthropic(api_key=api_key)

                        # Build dynamic prompt from SharePoint rules
                        prompt = build_dynamic_claude_prompt(ai_rules, full_text)

                        response = client.messages.create(
                            model="claude-3-haiku-20240307",
                            max_tokens=4096,
                            temperature=0.3,
                            messages=[{"role": "user", "content": prompt}]
                        )

                        response_text = response.content[0].text
                        json_start = response_text.find('{')
                        json_end = response_text.rfind('}') + 1

                        if json_start >= 0 and json_end > json_start:
                            result_json = json.loads(response_text[json_start:json_end])
                            corrected_text = result_json.get('corrected_text', '')
                            changes_count = result_json.get('changes_made', 0)

                            if changes_count > 0 and corrected_text:
                                # Apply corrections paragraph by paragraph
                                corrected_paras = corrected_text.split('\n\n')
                                para_index = 0

                                for para in doc.paragraphs:
                                    if para.text.strip() and para_index < len(corrected_paras):
                                        if len(para.runs) > 0:
                                            # Update first run, clear others
                                            para.runs[0].text = corrected_paras[para_index]
                                            for run in para.runs[1:]:
                                                run.text = ""
                                        para_index += 1

                                issues.append({
                                    'rule_name': 'AI Style Corrections',
                                    'rule_type': 'AI',
                                    'description': f"Found {changes_count} style violations",
                                    'location': 'Document-wide',
                                    'priority': 1
                                })
                                fixes_applied.append({
                                    'rule_name': 'AI Style Corrections',
                                    'rule_type': 'AI',
                                    'found_value': f'{changes_count} style violations',
                                    'fixed_value': 'British English, contractions, symbols corrected',
                                    'location': 'Document-wide'
                                })
                                logging.info(f"Claude style corrections: {changes_count}")
                            else:
                                logging.info("No spelling changes needed")

                    except Exception as e:
                        logging.error(f"Claude API error: {str(e)}")
                        issues.append({
                            'rule_name': 'AI Style Validation',
                            'rule_type': 'AI',
                            'description': f"AI style validation failed: {str(e)}",
                            'location': 'N/A',
                            'priority': 1
                        })
            else:
                logging.warning("ANTHROPIC_API_KEY not set - skipping AI style validation")

            # 2. Font checking (proven to work)
            font_changes = 0
            expected_font = "Arial"
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if run.text.strip():
                        current_font = run.font.name
                        if current_font is None or current_font != expected_font:
                            run.font.name = expected_font
                            font_changes += 1

            if font_changes > 0:
                fixes_applied.append({
                    'rule_name': 'All Text Font',
                    'rule_type': 'Font',
                    'found_value': f'{font_changes} runs with wrong font',
                    'fixed_value': expected_font,
                    'location': f'Document-wide ({font_changes} runs)'
                })
                logging.info(f"Font changes: {font_changes}")

            logging.info(f"v3.3 TOTAL: {len(issues)} issues, {len(fixes_applied)} fixes")

            result = {
                'document': doc,
                'issues': issues,
                'fixes_applied': fixes_applied
            }

            # Save fixed document
            logging.info('Saving fixed document to stream...')
            fixed_stream = BytesIO()
            result['document'].save(fixed_stream)
            fixed_stream.seek(0)

        elif file_extension in ['.vsdx', '.vsd']:
            result = validate_visio_document(file_stream, rules)

            # Save fixed document
            fixed_stream = BytesIO()
            result['document'].save_vsdx(fixed_stream)
            fixed_stream.seek(0)

        else:
            logging.error(f"Unsupported file type: {file_extension}")
            return func.HttpResponse(
                json.dumps({"error": f"Unsupported file type: {file_extension}"}),
                status_code=400
            )

        # 7. Upload fixed file (overwrite original) - only if file_url is provided
        logging.info("=" * 60)
        logging.info(f"UPLOAD CHECK v3.0:")
        logging.info(f"  file_url present: {file_url is not None}")
        logging.info(f"  fixes_applied count: {len(result.get('fixes_applied', []))}")
        logging.info(f"  Will upload: {file_url and result['fixes_applied']}")
        logging.info("=" * 60)

        if file_url and result['fixes_applied']:
            logging.info(f'Step 7: UPLOADING fixed document ({len(result["fixes_applied"])} fixes)...')
            upload_file(token, fixed_stream, file_url)
            logging.info('Step 7: Upload complete!')
        elif not file_url:
            logging.info('Step 7: Skipping upload (no file URL provided)')
        else:
            logging.info('Step 7: NO FIXES TO UPLOAD (fixes_applied is empty)')

        # 8. Generate report
        logging.info('Step 8: Generating report...')
        report_html = generate_report(file_name, result['issues'], result['fixes_applied'])
        report_url = None

        # Always upload report to SharePoint
        logging.info('Uploading report to SharePoint...')
        report_stream = BytesIO(report_html.encode('utf-8'))
        report_filename = f"{os.path.splitext(file_name)[0]}_ValidationReport.html"
        if file_url:
            report_folder = os.path.dirname(file_url)
            report_path = f"{report_folder}/{report_filename}" if report_folder else f"/{report_filename}"
        else:
            # No file_url — upload to a Validation Reports folder in the default drive
            report_path = f"/Validation Reports/{report_filename}"
        logging.info(f"Report will be uploaded to: {report_path}")
        try:
            report_url = upload_file(token, report_stream, report_path)
            logging.info(f"Report uploaded: {report_url}")
        except Exception as upload_err:
            logging.error(f"Failed to upload report: {str(upload_err)}")
            report_url = None

        # 8.5. Save validation result to Validation Results list
        logging.info('Step 8.5: Saving validation result to Validation Results list...')
        validation_result_info = None
        try:
            site_id = get_site_id(token)
            unfixed_issues = len(result['issues']) - len(result['fixes_applied'])
            result_status = "Passed" if unfixed_issues == 0 else "Failed"

            validation_result_info = save_validation_result(
                token=token,
                site_id=site_id,
                filename=file_name,
                issues_count=len(result['issues']),
                fixes_count=len(result['fixes_applied']),
                status=result_status,
                html_report=report_html,
                report_url=report_url
            )
            logging.info(f"✓ Validation result saved: {validation_result_info['list_item_url']}")

            # 8.6. Update document metadata with link to validation result
            if file_url and validation_result_info:
                logging.info('Step 8.6: Updating document with validation result link...')
                try:
                    update_success = update_document_metadata(
                        token=token,
                        site_id=site_id,
                        file_url=file_url,
                        validation_result_url=validation_result_info['list_item_url']
                    )
                    if update_success:
                        logging.info("✓ Document metadata updated with validation result link")
                    else:
                        logging.warning("Document metadata update returned False")
                except Exception as update_error:
                    logging.error(f"Failed to update document metadata: {str(update_error)}")
                    logging.error(f"Continuing with validation process...")
        except Exception as e:
            logging.error(f"Failed to save validation result: {str(e)}")
            logging.error(f"Continuing with validation process...")

        # 9. Update validation status
        logging.info('Step 9: Updating final validation status...')
        # Pass if no issues, or if all issues were auto-fixed
        unfixed_issues = len(result['issues']) - len(result['fixes_applied'])
        final_status = "Passed" if unfixed_issues == 0 else "Failed"
        logging.info(f"Issues: {len(result['issues'])}, Fixes: {len(result['fixes_applied'])}, Unfixed: {unfixed_issues}")
        update_validation_status(token, item_id, final_status, report_url)

        # 10. Return response with fixed file content
        logging.info(f'=== VALIDATION COMPLETE: {final_status} ===')

        response_data = {
            "status": final_status,
            "issuesFound": len(result['issues']),
            "issuesFixed": len(result['fixes_applied']),
            "reportUrl": report_url,
            "validationResultUrl": validation_result_info['list_item_url'] if validation_result_info else None,
            # Hyperlink objects for Power Automate (SharePoint format)
            "reportLink": {
                "Description": "View HTML Report",
                "Url": report_url
            } if report_url else None,
            "validationResultLink": {
                "Description": "View Validation Result",
                "Url": validation_result_info['list_item_url']
            } if validation_result_info else None
        }

        # Include fixed file content if fixes were applied
        if result['fixes_applied']:
            import base64
            fixed_stream.seek(0)
            fixed_content_base64 = base64.b64encode(fixed_stream.read()).decode('utf-8')
            response_data["fixedFileContent"] = fixed_content_base64
            logging.info(f"Returning fixed file content ({len(fixed_content_base64)} chars)")

        return func.HttpResponse(
            json.dumps(response_data),
            mimetype="application/json",
            status_code=200
        )

    except Exception as e:
        logging.error(f"=== VALIDATION FAILED ===")
        logging.error(f"Error: {str(e)}")
        logging.error(f"Error type: {type(e).__name__}")
        import traceback
        logging.error(f"Traceback: {traceback.format_exc()}")
        return func.HttpResponse(
            json.dumps({
                "error": str(e),
                "error_type": type(e).__name__
            }),
            mimetype="application/json",
            status_code=500
        )
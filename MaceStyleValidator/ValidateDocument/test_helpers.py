"""Helper functions for local testing"""
import json
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor

def create_test_word_doc():
    """Create a test Word doc with intentional issues"""
    doc = Document()
    
    # Add heading with wrong font
    heading = doc.add_heading('Test Heading', level=1)
    heading.runs[0].font.name = 'Times New Roman'  # Wrong font
    heading.runs[0].font.color.rgb = RGBColor(255, 0, 0)  # Wrong color
    
    # Add normal paragraph
    doc.add_paragraph('This is a test paragraph.')
    
    # Save to BytesIO
    stream = BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream

def create_mock_rules():
    """Create mock validation rules for testing"""
    return [
        {
            'title': 'Heading 1 Font',
            'rule_type': 'Font',
            'doc_type': 'Word',
            'check_value': 'Heading1Font',
            'expected_value': 'Arial',
            'auto_fix': True,
            'priority': 1
        },
        {
            'title': 'Heading 1 Color',
            'rule_type': 'Color',
            'doc_type': 'Word',
            'check_value': 'Heading1Color',
            'expected_value': '0,51,153',  # Mace blue
            'auto_fix': True,
            'priority': 2
        }
    ]

def create_test_request(file_name="test.docx"):
    """Create mock HTTP request for testing"""
    return {
        "itemId": 1,
        "fileUrl": "/sites/testsite/Shared Documents/test.docx",
        "fileName": file_name
    }
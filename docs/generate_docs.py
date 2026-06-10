"""Generate branded Word docs for the Mace Style Validator: User Guide/FAQ and Architecture/Security/Governance."""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
LOGO = os.path.join(HERE, "maceway-logo.png")
DIA_ARCH = os.path.join(HERE, "diagram-architecture.jpg")
DIA_PROCESS = os.path.join(HERE, "diagram-process.jpg")
DIA_PILOT = os.path.join(HERE, "diagram-pilot-prod.jpg")
GREY = RGBColor(0x79, 0x75, 0x74)
DARK = RGBColor(0x33, 0x33, 0x33)
GREEN = RGBColor(0x2E, 0x7D, 0x52)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)


def page_field(paragraph):
    run = paragraph.add_run()
    a = OxmlElement('w:fldChar'); a.set(qn('w:fldCharType'), 'begin')
    b = OxmlElement('w:instrText'); b.set(qn('xml:space'), 'preserve'); b.text = 'PAGE'
    c = OxmlElement('w:fldChar'); c.set(qn('w:fldCharType'), 'end')
    run._r.append(a); run._r.append(b); run._r.append(c)


def new_doc():
    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.8)
        s.left_margin = s.right_margin = Inches(0.9)
    n = doc.styles['Normal']; n.font.name = 'Arial'; n.font.size = Pt(11); n.font.color.rgb = DARK
    for name, sz in [('Heading 1', 16), ('Heading 2', 13), ('Heading 3', 11.5)]:
        st = doc.styles[name]; st.font.name = 'Arial'; st.font.size = Pt(sz)
        st.font.color.rgb = GREY; st.font.bold = True
    return doc


def footer(doc, label):
    p = doc.sections[0].footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{label}   ·   Confidential   ·   Page ")
    r.font.size = Pt(8); r.font.color.rgb = GREY; r.font.name = 'Arial'
    page_field(p)


def cover(doc, title, subtitle, doctype, version="1.0"):
    if os.path.exists(LOGO):
        doc.add_picture(LOGO, width=Inches(2.3))
    doc.add_paragraph()
    t = doc.add_table(rows=1, cols=1); t.autofit = True
    cell = t.rows[0].cells[0]; shade(cell, "797574")
    cp = cell.paragraphs[0]; cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = cp.add_run(title); r.font.size = Pt(26); r.font.bold = True; r.font.color.rgb = WHITE; r.font.name = 'Arial'
    sp = cell.add_paragraph(); rs = sp.add_run(subtitle)
    rs.font.size = Pt(13); rs.font.color.rgb = RGBColor(0xEC, 0xEC, 0xEC); rs.font.name = 'Arial'
    cell.add_paragraph()
    doc.add_paragraph()
    meta = doc.add_table(rows=4, cols=2)
    for i, (k, v) in enumerate([
        ("Document", doctype), ("Version", version),
        ("Owner", "Stephen Cummins  ·  stephen.cummins@macegroup.com"),
        ("Date", "June 2026"),
    ]):
        kk = meta.rows[i].cells[0].paragraphs[0].add_run(k)
        kk.font.bold = True; kk.font.size = Pt(10); kk.font.color.rgb = GREY; kk.font.name = 'Arial'
        vv = meta.rows[i].cells[1].paragraphs[0].add_run(v)
        vv.font.size = Pt(10); vv.font.name = 'Arial'
    doc.add_page_break()


def body(doc, text, bold=False):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.font.bold = bold
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
    p.add_run(text)
    return p


def kv_table(doc, rows, col0="Item", col1="Detail"):
    t = doc.add_table(rows=1, cols=2); t.style = 'Light Grid Accent 1'
    hdr = t.rows[0].cells
    for c, txt in zip(hdr, (col0, col1)):
        run = c.paragraphs[0].add_run(txt); run.font.bold = True; run.font.size = Pt(10); run.font.name = 'Arial'
        shade(c, "797574"); run.font.color.rgb = WHITE
    for k, v in rows:
        cells = t.add_row().cells
        rk = cells[0].paragraphs[0].add_run(k); rk.font.bold = True; rk.font.size = Pt(10); rk.font.name = 'Arial'
        rv = cells[1].paragraphs[0].add_run(v); rv.font.size = Pt(10); rv.font.name = 'Arial'
    doc.add_paragraph()
    return t


def image(doc, path, width=Inches(6.5), caption=None):
    if not os.path.exists(path):
        return
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=width)
    if caption:
        cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cp.add_run(caption); r.italic = True; r.font.size = Pt(9)
        r.font.color.rgb = GREY; r.font.name = 'Arial'
    doc.add_paragraph()


def faq(doc, q, a):
    p = doc.add_paragraph(); r = p.add_run("Q.  " + q); r.font.bold = True; r.font.color.rgb = GREY
    pa = doc.add_paragraph(); pa.add_run("A.  " + a)
    pa.paragraph_format.space_after = Pt(10)


# ============================================================= DOC 1: USER GUIDE / FAQ
d = new_doc()
cover(d, "Mace Style Validator", "User Guide & Frequently Asked Questions",
      "User Guide & FAQ")

d.add_heading("1. Overview", level=1)
body(d, "The Mace Style Validator automatically checks documents in the Mace Way Control Centre against the "
        "Control Centre Writing Style Guide, fixes common issues, and produces a branded validation report — "
        "all from within SharePoint. It works across Word, Excel, PowerPoint and Visio.")

d.add_heading("2. Running a validation", level=1)
body(d, "Validation is triggered from the document library — there is no separate application to open.")
for t in [
    'Open the Document Technical Review library and locate your document.',
    'Set the document\'s ValidationStatus column to "Validate Now".',
    'The validator runs automatically: it checks the document against the live style rules, auto-fixes common '
    'issues, generates a report, and records the outcome.',
    'Within a short time the document\'s status updates (Passed, Review Required or Failed) and a "View Report" '
    'link appears.',
]:
    bullet(d, t)

image(d, DIA_PROCESS, caption="The end-to-end validation process")

d.add_heading("3. Understanding the report", level=1)
body(d, "Each validation produces a Mace-branded HTML report containing:")
for t in [
    "A status banner — Passed, Review Required or Failed.",
    "A summary — issues found, auto-fixed, and remaining.",
    "Changes Made — each rule applied, what was found and what it was changed to, and where.",
    "Remaining Issues — anything needing manual attention.",
    "Links back to the source document and its library.",
]:
    bullet(d, t)
body(d, "Tip: open report links with a normal click — they navigate in the same tab. Use your browser's Back "
        "button to return to the report.")

d.add_heading("4. Where results are recorded", level=1)
kv_table(d, [
    ("ValidationStatus", "Passed / Review Required / Failed, shown on the document."),
    ("LastValidated", "Date and time of the most recent run."),
    ("ValidationReport", '"View Report" — opens the HTML report for the document.'),
    ("ValidationResultLink", '"View Result" — opens the matching record in the Validation Results list.'),
    ("Validation Results list", "One record per run: file name, date, status, issues found and fixed."),
    ("Validation Reports library", "Holds the generated HTML reports."),
], col0="Where", col1="What it shows")

d.add_heading("5. What it checks", level=1)
for t in [
    "British English spelling (e.g. colour, analyse, centre).",
    "Contraction expansion and symbol replacement (& → and, % → percent).",
    "Number formatting and font standardisation (Arial).",
    "Document structure and headings.",
]:
    bullet(d, t)
body(d, "Supported formats: Word (.docx), Excel (.xlsx), PowerPoint (.pptx) and Visio (.vsdx).")

d.add_heading("6. Managing the style rules", level=1)
body(d, "The rules are not hard-coded — they live in the Style Rules list in SharePoint and are owned by the "
        "business. To change what the validator enforces, edit that list; no software release is required. "
        "Add a new item to introduce a rule, or amend an existing one to change its behaviour.")

d.add_heading("7. Frequently asked questions", level=1)
faq(d, "Does the validator send my documents outside Mace?",
    "No. All documents, rules, results and reports stay within the Mace tenant. The validation engine processes "
    "content in memory and stores nothing. External AI is currently disabled, so no document content is sent to "
    "any third-party AI service.")
faq(d, "Will it change my original document?",
    "It corrects common, auto-fixable issues and records exactly what changed in the report. Anything that needs "
    "judgement is listed under Remaining Issues for you to address manually.")
faq(d, "What do the statuses mean?",
    "Passed — compliant (any issues were auto-fixed). Review Required — some issues were fixed but others need "
    "manual attention. Failed — issues were found that could not be auto-fixed.")
faq(d, "Why does the report sometimes show a status of 'Validating…'?",
    "That is a transient state while a run is in progress; it updates to the final status when the run completes.")
faq(d, "Who can see the reports and results?",
    "Anyone with access to the Mace Way Control Centre library and lists — the same permissions as the documents "
    "themselves. Access is governed entirely by Mace.")
faq(d, "How do I add or change a style rule?",
    "Edit the Style Rules list in SharePoint. Changes take effect on the next validation — no deployment needed.")
faq(d, "The logo doesn't appear if I download the report. Is that expected?",
    "Yes. Reports are designed to be viewed from SharePoint, where the Mace logo loads from the site. A downloaded "
    "copy may not show it.")
faq(d, "Who do I contact for help?",
    "Stephen Cummins — stephen.cummins@macegroup.com.")
footer(d, "Mace Style Validator — User Guide & FAQ")
d.save(os.path.join(HERE, "MaceStyleValidator-User-Guide-and-FAQ.docx"))
print("wrote User Guide & FAQ")


# ============================================================= DOC 2: ARCH / SECURITY / GOVERNANCE
d = new_doc()
cover(d, "Mace Style Validator", "Architecture, Security, Deployment & Governance",
      "Architecture, Security & Governance")

d.add_heading("1. Executive summary", level=1)
body(d, "The Mace Style Validator enforces the Mace Way Control Centre Writing Style Guide automatically, at the "
        "point documents are authored. It is integrated into SharePoint and orchestrated by Power Automate, with a "
        "lightweight, stateless validation engine running as an Azure Function. This document sets out the "
        "architecture, the security and data-residency model, the deployment approach, and the governance "
        "recommendations for taking the capability to production.")
body(d, "Key message on security: although the validation engine currently runs in a separate Azure subscription, "
        "all data and all access controls remain within the Mace tenant. The engine stores nothing, sees only one "
        "SharePoint site, and can be revoked by Mace at any time.", bold=True)

d.add_heading("2. Architecture", level=1)
body(d, "The solution has four components:")
for t in [
    "Mace SharePoint (Mace Way Control Centre) — holds the documents, the style rules, the validation results and "
    "the generated reports. This is the system of record and never leaves the Mace tenant.",
    "Power Automate flow — triggered when a document is flagged for validation. It runs inside Mace, as a Mace "
    "user, and performs all writes back to SharePoint.",
    "Azure Function (validation engine) — receives a document, validates it against the rules, returns the result "
    "and a report. It is stateless and holds no data.",
    "Microsoft Entra app registration — the identity the engine uses to read style rules, scoped to a single site "
    "via Sites.Selected.",
]:
    bullet(d, t)
image(d, DIA_ARCH, caption="Figure 1 — Architecture and trust boundary: the engine is separate, but data stays in Mace")
body(d, "Flow of a validation: the flow sends the document content and reads the rules; the engine validates in "
        "memory and returns the outcome plus a report; the flow writes the results, uploads the report, and "
        "updates the document — all within SharePoint.")
image(d, DIA_PROCESS, caption="Figure 2 — End-to-end validation process")

d.add_heading("3. Security & data residency", level=1)
d.add_heading("3.1 Your data stays in your tenant", level=2)
for t in [
    "All documents, rules, results and reports live in Mace SharePoint. The engine never persists them elsewhere.",
    "The engine is stateless — content is processed in memory and discarded. There is no database and no retention "
    "in the external subscription.",
    "All traffic is encrypted in transit (TLS / HTTPS).",
    "External AI is currently disabled; no document content is sent to any third-party AI service. If enabled in "
    "future, a data-classification check warns before any external processing.",
]:
    bullet(d, t)
d.add_heading("3.2 Identity & access", level=2)
for t in [
    "Access is via an Entra app registration on the Mace tenant, admin-consented by Mace IT.",
    "It uses Sites.Selected — scoped to the Mace Way Control Centre site only, not tenant-wide. It cannot reach "
    "any other Mace data.",
    "The engine's HTTP endpoints require a function key; only the Mace flow can invoke them.",
    "Mace can revoke access instantly — withdraw the site grant, disable the app registration, or rotate the key.",
]:
    bullet(d, t)
d.add_heading("3.3 Auditing", level=2)
for t in [
    "Every request is assigned a request identifier and emits an audit event.",
    "Access-control checks are enforced on each call.",
    "Controls are aligned to recognised SOC 2 control families (access control, monitoring).",
]:
    bullet(d, t)

d.add_heading("4. Why the engine runs where it does", level=1)
body(d, "The engine currently runs in a personal Azure subscription. This was a deliberate pilot decision: it "
        "enabled rapid iteration with no dependency on Mace infrastructure or change-control while the concept was "
        "proven. The security model described above does not depend on where the engine runs, because the data and "
        "the controls live in Mace. The engine is tenant-agnostic: moving it to a Mace-owned subscription is a "
        "redeployment of the same code, with no change to the data model or security posture.")

image(d, DIA_PILOT, caption="Figure 3 — From pilot to production: same code, same security model")

d.add_heading("5. Deployment & operations", level=1)
kv_table(d, [
    ("Validation engine", "Azure Function, Python, serverless (Flex Consumption)."),
    ("Orchestration", "Power Automate flow (Mace Document Style Validator)."),
    ("Storage", "SharePoint lists (Style Rules, Validation Results) and libraries (documents, Validation Reports)."),
    ("Infrastructure-as-code", "ARM template defines the Azure resources."),
    ("CI/CD", "Azure DevOps pipeline for build and deploy."),
    ("Configuration", "Style rules and results are managed in SharePoint — no release needed to change a rule."),
    ("Current state", "Piloted on Mace Way Control Centre; validated end-to-end on production documents."),
], col0="Aspect", col1="Detail")
body(d, "Operational changes fall into two categories: rule changes (made by the business in SharePoint, no "
        "deployment) and engine changes (deployed via the pipeline). Access can be revoked by Mace IT independently "
        "of either.")

d.add_heading("6. Governance & recommendations", level=1)
for t in [
    "Adopt the engine into a Mace-owned Azure subscription under CDIO governance, with a nominated DTS owner for "
    "the app registration and hosting.",
    "Tighten the site grant to least-privilege: the engine only needs read access to the style rules, so the grant "
    "can be reduced from its current level to read-only.",
    "Maintain the app registration, function key and site grant under standard Mace identity and secrets "
    "management.",
    "Keep style-rule ownership with the business; keep engine changes under change control.",
]:
    bullet(d, t)

d.add_heading("7. Risks & mitigations", level=1)
kv_table(d, [
    ("Engine hosted outside Mace (pilot)", "Mitigated: no data stored outside; least-privilege, revocable access. "
                                           "Resolved by migrating to Mace Azure."),
    ("Over-privileged site grant", "Reduce to read-only; tracked as a recommendation."),
    ("Credential management", "Standard secret rotation; no secrets in code or documents."),
    ("Service availability", "Serverless platform with retry in the flow; non-critical, asynchronous workload."),
], col0="Risk", col1="Mitigation")

d.add_heading("8. Configuration summary", level=1)
kv_table(d, [
    ("App registration", "MaceStyleValidator-App (Mace tenant)."),
    ("Permission model", "Sites.Selected — single site (Mace Way Control Centre)."),
    ("SharePoint site", "mace365.sharepoint.com/sites/MaceWayControlCentre."),
    ("Validation engine", "Azure Function (Python, Flex Consumption), UK South."),
    ("External AI", "Disabled."),
], col0="Item", col1="Value")
body(d, "Note: this document contains no secrets. Client secrets and keys are held in Azure / Mace secret "
        "management and are never recorded here.")
footer(d, "Mace Style Validator — Architecture, Security & Governance")
d.save(os.path.join(HERE, "MaceStyleValidator-Architecture-Security-Governance.docx"))
print("wrote Architecture/Security/Governance")

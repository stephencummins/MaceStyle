#!/usr/bin/env python3
"""Generate the editable Word version of the Mace funding brief (AI Service Lead Brief).

Mirrors mace-funding-brief.html/.pdf but as an editable .docx Tobias can annotate.
Run: python3 generate_brief_docx.py  ->  mace-funding-brief.docx
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Mace brand palette
INK = RGBColor(0x13, 0x0B, 0x01)
TEAL = RGBColor(0x00, 0x92, 0x9F)
PURPLE = RGBColor(0x89, 0x2B, 0x90)
GOLD = RGBColor(0xFD, 0xB9, 0x24)
MUTE = RGBColor(0x5B, 0x58, 0x50)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BODY_FONT = "Arial"          # Mace house body font
HEAD_FONT = "Georgia"        # serif display, stands in for Fraunces

doc = Document()
doc.core_properties.title = "MaceStyle — Bringing It In-House"
doc.core_properties.author = "Stephen Cummins"

# Base style
normal = doc.styles["Normal"]
normal.font.name = BODY_FONT
normal.font.size = Pt(10.5)
normal.font.color.rgb = INK
for section in doc.sections:
    section.top_margin = section.bottom_margin = Cm(1.8)
    section.left_margin = section.right_margin = Cm(2.0)


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hexcolor)
    tcPr.append(sh)


def para(text="", size=10.5, color=INK, bold=False, italic=False, font=BODY_FONT,
         after=6, before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    if align:
        p.alignment = align
    r = p.add_run(text)
    r.font.name, r.font.size, r.font.bold, r.font.italic = font, Pt(size), bold, italic
    r.font.color.rgb = color
    return p


def rich(segments, size=10.5, after=6, before=0):
    """segments: list of (text, {bold,italic,color,font})."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    for text, opt in segments:
        r = p.add_run(text)
        r.font.name = opt.get("font", BODY_FONT)
        r.font.size = Pt(size)
        r.font.bold = opt.get("bold", False)
        r.font.italic = opt.get("italic", False)
        r.font.color.rgb = opt.get("color", INK)
    return p


def heading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(4)
    bar = p.add_run("▌ ")   # left block as a teal accent
    bar.font.color.rgb = TEAL
    bar.font.size = Pt(14)
    r = p.add_run(text)
    r.font.name, r.font.size, r.font.bold = HEAD_FONT, Pt(14), True
    r.font.color.rgb = INK


# --- Header ---
para("AI SERVICE LEAD BRIEF", size=8.5, color=TEAL, bold=True, after=2)
para("MaceStyle: Bringing It In-House", size=24, color=INK, bold=True, font=HEAD_FONT, after=2)
para("Running our Style Guide validator on Mace's own billing and governance",
     size=13, color=PURPLE, italic=True, font=HEAD_FONT, after=8)
rich([("Prepared for ", {"color": MUTE}), ("Tobias", {"bold": True, "color": MUTE}),
      (" by Stephen Cummins · AI Service Lead, Mace Digital · 2 July 2026 (updated 5 July 2026)", {"color": MUTE})],
     size=8.5, after=10)

# --- Lead callout (single-cell shaded table w/ gold left border) ---
lead_tbl = doc.add_table(rows=1, cols=1)
lead_tbl.autofit = True
cell = lead_tbl.rows[0].cells[0]
shade(cell, "F6F8F8")
cell.paragraphs[0].paragraph_format.space_after = Pt(0)
lr = cell.paragraphs[0]
segs = [
    ("MaceStyle now runs its AI on ", {}),
    ("GPT via Azure OpenAI", {"bold": True, "color": PURPLE}),
    (" — ", {}), ("Microsoft-native", {"bold": True, "color": PURPLE}),
    (", on Azure billing and governance, at ", {}), ("pennies per document", {"bold": True, "color": PURPLE}),
    (". I've already made the switch and proven it end-to-end on the live pilot. To make it a Mace service "
     "it needs a Mace Azure subscription to own it and a data-governance nod. Then it's roughly half a day "
     "to cut over.", {}),
]
for text, opt in segs:
    r = lr.add_run(text)
    r.font.name, r.font.size = BODY_FONT, Pt(11)
    r.font.bold = opt.get("bold", False)
    r.font.color.rgb = opt.get("color", INK)
para("", after=2)

# --- What we have today ---
heading("What we have today")
para("MaceStyle reads documents in the Mace Way Control Centre and checks them against the Writing "
     "Style Guide — British spelling, contractions, symbols, number and font standards — flagging "
     "issues and proposing corrections. It is live in pilot, validated by testers (Natasha, Jade) in June. "
     "The intelligent rules — the judgement calls a simple find-and-replace can't make — are powered by AI.")
rich([("What changed this month: ", {"bold": True}),
      ("I re-architected MaceStyle so the AI backend is ", {}), ("swappable", {"italic": True}),
      (" via a single setting, and switched the live pilot onto ", {}),
      ("GPT-5 (Azure OpenAI)", {"bold": True, "color": PURPLE}),
      (" — Microsoft's own AI service, running inside Azure. It's proven working end-to-end. The same code "
       "can run Claude instead with no rewrite, so we are not locked to one model.", {})])
rich([("The catch: ", {"bold": True}),
      ("it currently runs on ", {}), ("my personal", {"italic": True}),
      (" Azure subscription. That proves the concept — it's not how Mace should run a service it depends "
       "on. To make this a Mace service, the billing and governance need to sit inside Mace.", {})])

# --- Proposal + table ---
heading("The proposal: run it on Mace's Azure — Microsoft-native")
rich([("Mace's instinct is to stay on Microsoft wherever possible. This fits that exactly. The AI now runs "
       "on ", {}), ("Azure OpenAI", {"bold": True, "color": PURPLE}),
      (" — the GPT models (the ChatGPT family) delivered as a first-party Microsoft Azure service:", {})])

rows = [
    ("", "Today (my Azure)", "Proposed (Mace's Azure)"),
    ("Who pays", "Me, personally", "Mace's existing Azure invoice"),
    ("What it is", "GPT via Azure OpenAI", "Same — a native Azure service Mace already has"),
    ("Governance", "Mine", "Entra ID, RBAC, data zone, spend caps"),
    ("Procurement", "—", "None. No new vendor, no third-party sign-up"),
]
tbl = doc.add_table(rows=len(rows), cols=3)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = "Table Grid"
for ci, text in enumerate(rows[0]):
    c = tbl.rows[0].cells[ci]
    shade(c, "130B01")
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text.upper())
    r.font.name, r.font.size, r.font.bold = BODY_FONT, Pt(8.5), True
    r.font.color.rgb = WHITE
for ri in range(1, len(rows)):
    for ci, text in enumerate(rows[ri]):
        c = tbl.rows[ri].cells[ci]
        if ri % 2 == 0:
            shade(c, "F6F8F8")
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        r.font.name, r.font.size = BODY_FONT, Pt(9.5)
        r.font.bold = (ci == 0) or (ci == 2)
        if ci == 0:
            r.font.color.rgb = TEAL
        elif ci == 2:
            r.font.color.rgb = PURPLE
para("", after=2)
rich([("The key point: ", {"bold": True}),
      ("there is no new supplier and no AI contract to negotiate.", {"bold": True}),
      (" Azure OpenAI is Microsoft, on the Azure agreement Mace already runs. Standing it up in a Mace "
       "subscription is deploying a model and pointing the app at it — no marketplace step, no external "
       "account.", {})])
para("", after=2)
rich([("Model choice stays open. ", {"bold": True, "color": PURPLE}),
      ("Because the backend is swappable, Mace can run ", {}), ("Claude", {"bold": True}),
      (" instead (via Microsoft Foundry, also inside Azure) for content that benefits from its slightly "
       "more careful editing — a one-setting change, not a rebuild. My recommendation: ", {}),
      ("default to GPT / Azure OpenAI", {"bold": True}),
      (" (easiest to provision and the cheaper of the two) and keep Claude available as an option.", {})])

# --- Architecture (text flow, editable) ---
flow = doc.add_paragraph()
flow.alignment = WD_ALIGN_PARAGRAPH.CENTER
flow.paragraph_format.space_before = Pt(4)
for i, (label, col) in enumerate([("Mace Azure", INK), ("MaceStyle Function", TEAL),
                                  ("Azure OpenAI", PURPLE), ("GPT-5", GOLD)]):
    if i:
        arw = flow.add_run("  →  ")
        arw.font.color.rgb = TEAL
        arw.font.bold = True
    r = flow.add_run(label)
    r.font.name, r.font.size, r.font.bold = BODY_FONT, Pt(10), True
    r.font.color.rgb = col
para("Everything stays inside Mace's Azure agreement — a first-party Microsoft service, no third party.",
     size=8, color=MUTE, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)

# --- Cost ---
heading("What it costs")
para("Consumption-based and genuinely cheap. I measured both models on real documents through the live "
     "pipeline:", after=6)
cost_rows = [
    ("Document size", "GPT-5-mini (recommended)", "Claude (alternative)"),
    ("~1 page", "£0.003", "£0.004"),
    ("~3 pages", "£0.004", "£0.008"),
    ("~8 pages", "£0.008", "£0.015"),
]
ctbl = doc.add_table(rows=len(cost_rows), cols=3)
ctbl.alignment = WD_TABLE_ALIGNMENT.CENTER
ctbl.style = "Table Grid"
for ci, text in enumerate(cost_rows[0]):
    c = ctbl.rows[0].cells[ci]
    shade(c, "130B01")
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text.upper())
    r.font.name, r.font.size, r.font.bold = BODY_FONT, Pt(8.5), True
    r.font.color.rgb = WHITE
for ri in range(1, len(cost_rows)):
    for ci, text in enumerate(cost_rows[ri]):
        c = ctbl.rows[ri].cells[ci]
        if ri % 2 == 0:
            shade(c, "F6F8F8")
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        r.font.name, r.font.size = BODY_FONT, Pt(9.5)
        r.font.bold = (ci == 1)
        if ci == 1:
            r.font.color.rgb = PURPLE
para("US-dollar list prices, converted approximately; Mace's enterprise Azure agreement would likely "
     "discount further.", size=8, color=MUTE, italic=True, after=6)
rich([("GPT is ", {}), ("~35–50% cheaper per document", {"bold": True}),
      (" than Claude, and the gap widens with length — but in absolute terms both are a ", {}),
      ("rounding error", {"bold": True}),
      (": realistic Control Centre volume is a few pounds a month. No upfront commitment; Mace pays only "
       "for what it uses. Cost isn't the deciding factor — ease of ownership is, and Azure OpenAI wins on "
       "both.", {})])

# --- What I need from you ---
heading("What I need from you")
asks = [
    ("A home for it", " — nominate a Mace Azure subscription and cost centre to own the service. Mace "
     "already runs on Azure (the mace365 tenant, Power Platform), so nothing new is stood up."),
    ("A governance nod", " — document text is processed by a US-hosted model. Azure OpenAI's data zone "
     "and Azure-native governance are built for exactly this; it needs Sapna's sign-off, which I'll tee up."),
    ("Half a day of my time", " — to cut over. I've already built and proven the whole path on the live "
     "pilot, so this is execution, not discovery."),
]
for i, (lead, rest) in enumerate(asks, 1):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(lead)
    r.font.name, r.font.size, r.font.bold = BODY_FONT, Pt(10.5), True
    r2 = p.add_run(rest)
    r2.font.name, r2.font.size = BODY_FONT, Pt(10.5)

# --- Closing (dark shaded box) ---
close_tbl = doc.add_table(rows=1, cols=1)
cc = close_tbl.rows[0].cells[0]
shade(cc, "130B01")
p0 = cc.paragraphs[0]
p0.paragraph_format.space_after = Pt(4)
r = p0.add_run("Why this matters beyond MaceStyle")
r.font.name, r.font.size, r.font.bold = HEAD_FONT, Pt(13), True
r.font.color.rgb = WHITE
def close_para(segments):
    p = cc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    for text, opt in segments:
        r = p.add_run(text)
        r.font.name, r.font.size = BODY_FONT, Pt(10)
        r.font.bold = opt.get("bold", False)
        r.font.color.rgb = opt.get("color", RGBColor(0xD9, 0xD5, 0xCC))
close_para([("MaceStyle is the ", {}), ("proof, not the point.", {"bold": True, "color": GOLD}),
            (" It shows a repeatable pattern: identify a real Mace problem, stand up an AI service to "
             "solve it, and hand it over running on Mace's own Microsoft/Azure billing and governance — "
             "safely, cheaply, and without a procurement cycle.", {})])
close_para([("That pattern is the AI Service Lead role in practice. Give me the mandate and MaceStyle "
             "becomes the ", {}), ("first of several", {"bold": True, "color": GOLD}),
            (" — each one an internal AI service that pays for itself in hours saved and lands inside "
             "Mace's existing Azure controls from day one.", {})])

para("", after=2)
para("Technical detail for whoever provisions it: see the companion Handover Runbook (covers both the "
     "Azure OpenAI / GPT and Foundry / Claude setups).  ·  "
     "Stephen Cummins · AI Service Lead, Mace Digital · stephen.cummins@macegroup.com",
     size=8, color=MUTE, italic=True)

out = Path(__file__).parent / "mace-funding-brief.docx"
doc.save(str(out))
print("Wrote", out)
